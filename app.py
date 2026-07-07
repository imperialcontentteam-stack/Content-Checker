"""
SLC Course Content Checker — v4 (Auth + Filters + Spec Docs + Validation Reports)
=================================================================================
Internal tool for the South London College course content team.

What's new in v4
----------------
1.  User authentication — separate Admin and User accounts (secure PBKDF2 hashes)
    · Admin: import tracker, upload/manage spec documents, manage courses & users
    · User:  Run Check page only — run validation, view & download reports
2.  Import Courses reads Category ID (Number), Level and Type from the tracker
    and populates the Run Check filters automatically (dynamic, from the DB)
3.  Spec Extraction — upload qualification specification documents (PDF/DOCX)
    or extract from URL; Entry Requirements / Qualification Specification /
    Method of Assessment sections are auto-extracted and editable
4.  Run Check — Category ID / Level / Type filters, course details on the left,
    specification + extracted requirements on the right, then validation of:
    Qualification Specification · Entry Requirements · Method of Assessment (wording)
5.  Validation report in the agreed layout: red section boxes, numbered
    "Errors identified" (red) + "Recommend Action" (green), a green **No Errors**
    badge when the course is clean, and an issue summary
6.  Download Report on the same page — PDF and Word (.docx)
7.  The existing Content / Grammar / Other checks are kept untouched as tabs
    inside Run Check (Content Check · Grammar Check · Other Checks)

Stack: Python · Streamlit · SQLite · OpenRouter API · reportlab
Run:   streamlit run app.py
Default logins (change them in 👥 Users):  admin / admin123   ·   user / user123
"""

import hashlib
import html
import io
import json
import os
import re
import secrets as pysecrets
import sqlite3
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from difflib import SequenceMatcher

import pandas as pd
import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# ═══════════════════════════════════════════════════════════════════
#  CONSTANTS
# ═══════════════════════════════════════════════════════════════════

DB_PATH = "slc_checker.db"

FIELD_OPTIONS = {
    "entry_requirements": "Entry Requirements",
    "method_of_assessment": "Method of Assessment",
    "course_overview": "Course Overview",
}

# The three sections compared by the new validation (Run Check → Content Check)
VALIDATION_SECTIONS = [
    # (key on course row [tracker side], key on spec side, display label)
    ("qualification_spec_current", "spec_qualification", "Qualification Specification"),
    ("entry_requirements", "spec_entry_requirements", "Entry Requirement"),
    ("method_of_assessment", "spec_assessment", "Method of Assessment"),
]

COURSE_TYPES = ["Award", "Certificate", "Diploma"]

MODEL = "deepseek/deepseek-v4-pro"

RED = "#D71920"
GREEN = "#1E9E3E"

# Colour system for the Quality Review markup (mirrors the reference screenshots)
QR_CATEGORIES = {
    "grammar":            {"label": "Grammar",            "bg": "#FDD8D6", "border": "#E5484D"},
    "article":            {"label": "Articles (a/an/the)","bg": "#FFE3C7", "border": "#F76B15"},
    "spelling":           {"label": "Spelling",           "bg": "#E9DDFB", "border": "#8E4EC6"},
    "punctuation":        {"label": "Punctuation & Commas","bg": "#D5E7FB", "border": "#0090FF"},
    "capitalisation":     {"label": "Capitalisation",     "bg": "#D8F3DE", "border": "#30A46C"},
    "proper_noun":        {"label": "Proper Nouns",       "bg": "#FBDCEF", "border": "#D6409F"},
    "sentence_structure": {"label": "Sentence Structure", "bg": "#FBE8B4", "border": "#B58A00"},
    "consistency":        {"label": "Consistency",        "bg": "#D9F0F4", "border": "#0894B3"},
}

USER_AGENT = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
}

# ═══════════════════════════════════════════════════════════════════
#  DATABASE
# ═══════════════════════════════════════════════════════════════════

def get_conn():
    # timeout=30 → wait up to 30s for a lock instead of raising
    # "database is locked" when parallel workers write concurrently
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with get_conn() as c:
        c.executescript(
            """
            CREATE TABLE IF NOT EXISTS courses (
                id                   INTEGER PRIMARY KEY AUTOINCREMENT,
                course_name          TEXT UNIQUE NOT NULL,
                course_url           TEXT,
                spec_url             TEXT,
                entry_requirements   TEXT,
                method_of_assessment TEXT,
                course_overview      TEXT,
                spec_text            TEXT,
                updated_at           TEXT
            );

            CREATE TABLE IF NOT EXISTS reports (
                id             INTEGER PRIMARY KEY AUTOINCREMENT,
                course_id      INTEGER NOT NULL REFERENCES courses(id) ON DELETE CASCADE,
                checked_at     TEXT,
                fields_checked TEXT,
                status         TEXT,           -- 'pass' | 'errors' | 'failed'
                errors_json    TEXT,
                rewrites_json  TEXT,
                summary        TEXT
            );

            CREATE TABLE IF NOT EXISTS users (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                username      TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                salt          TEXT NOT NULL,
                role          TEXT NOT NULL DEFAULT 'user',   -- 'admin' | 'user'
                created_at    TEXT
            );

            CREATE TABLE IF NOT EXISTS validation_reports (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                course_id    INTEGER NOT NULL REFERENCES courses(id) ON DELETE CASCADE,
                checked_at   TEXT,
                checked_by   TEXT,
                status       TEXT,             -- 'pass' | 'errors' | 'failed'
                summary      TEXT,
                results_json TEXT
            );

            -- One row per UNIQUE qualification specification document.
            -- Several courses can point at the same document; extraction
            -- happens ONCE per document and is reused for every course.
            CREATE TABLE IF NOT EXISTS spec_documents (
                id             INTEGER PRIMARY KEY AUTOINCREMENT,
                spec_url       TEXT UNIQUE,
                filename       TEXT,
                raw_text       TEXT,
                extracted_json TEXT,            -- structured AI extraction
                status         TEXT DEFAULT 'pending',  -- 'pending'|'processed'|'error'
                method         TEXT,            -- 'ai' | 'heuristic'
                error          TEXT,
                content_hash   TEXT,
                processed_at   TEXT,
                updated_at     TEXT
            );
            """
        )
        # ── migrations ──────────────────────────────────────────────
        cols = [r[1] for r in c.execute("PRAGMA table_info(reports)")]
        if "rewrites_json" not in cols:
            c.execute("ALTER TABLE reports ADD COLUMN rewrites_json TEXT")

        course_cols = [r[1] for r in c.execute("PRAGMA table_info(courses)")]
        for col in ["category_id", "level", "course_type", "qualification_spec_current",
                    "spec_filename", "spec_entry_requirements",
                    "spec_qualification", "spec_assessment"]:
            if col not in course_cols:
                c.execute(f"ALTER TABLE courses ADD COLUMN {col} TEXT")
        if "spec_doc_id" not in course_cols:
            c.execute("ALTER TABLE courses ADD COLUMN spec_doc_id INTEGER "
                      "REFERENCES spec_documents(id)")


# ── authentication helpers ──────────────────────────────────────────

def _hash_password(password: str, salt: str) -> str:
    return hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"),
                               salt.encode("utf-8"), 200_000).hex()


def create_user(username: str, password: str, role: str = "user") -> bool:
    salt = pysecrets.token_hex(16)
    try:
        with get_conn() as c:
            c.execute(
                "INSERT INTO users (username, password_hash, salt, role, created_at) "
                "VALUES (?,?,?,?,?)",
                (username.strip(), _hash_password(password, salt), salt, role,
                 datetime.now().isoformat(timespec="seconds")),
            )
        return True
    except sqlite3.IntegrityError:
        return False


def set_password(user_id: int, password: str):
    salt = pysecrets.token_hex(16)
    with get_conn() as c:
        c.execute("UPDATE users SET password_hash=?, salt=? WHERE id=?",
                  (_hash_password(password, salt), salt, user_id))


def delete_user(user_id: int):
    with get_conn() as c:
        c.execute("DELETE FROM users WHERE id=?", (user_id,))


def all_users() -> list:
    with get_conn() as c:
        return [dict(r) for r in c.execute(
            "SELECT id, username, role, created_at FROM users ORDER BY role, username")]


def verify_login(username: str, password: str):
    """Return the user row (dict) if the credentials are valid, else None."""
    with get_conn() as c:
        row = c.execute("SELECT * FROM users WHERE username=?",
                        (username.strip(),)).fetchone()
    if not row:
        return None
    if pysecrets.compare_digest(row["password_hash"],
                                _hash_password(password, row["salt"])):
        return {"id": row["id"], "username": row["username"], "role": row["role"]}
    return None


def ensure_default_users():
    """Seed a default admin + user account on first run."""
    with get_conn() as c:
        n = c.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    if n == 0:
        create_user("admin", "admin123", "admin")
        create_user("user", "user123", "user")


# ── course helpers ──────────────────────────────────────────────────

def parse_level_type(course_name: str):
    """Derive Level and Type (Award/Certificate/Diploma) from a course name."""
    name = str(course_name or "")
    m_lvl = re.search(r"level\s*(\d+)", name, re.I)
    m_typ = re.search(r"\b(award|certificate|diploma)\b", name, re.I)
    return (m_lvl.group(1) if m_lvl else None,
            m_typ.group(1).title() if m_typ else None)


def upsert_course(row: dict) -> str:
    """Insert or update a course by name. Returns 'inserted' or 'updated'."""
    now = datetime.now().isoformat(timespec="seconds")

    # auto-derive Level / Type from the name when the tracker doesn't provide them
    lvl, typ = parse_level_type(row["course_name"])
    row.setdefault("level", None)
    row.setdefault("course_type", None)
    level = row.get("level") or lvl
    ctype = row.get("course_type") or typ

    with get_conn() as c:
        cur = c.execute("SELECT id FROM courses WHERE course_name = ?", (row["course_name"],))
        existing = cur.fetchone()
        if existing:
            c.execute(
                """UPDATE courses SET course_url=?, spec_url=?, entry_requirements=?,
                   method_of_assessment=?, course_overview=?, category_id=?,
                   level=?, course_type=?, qualification_spec_current=?, updated_at=?
                   WHERE id=?""",
                (row.get("course_url"), row.get("spec_url"), row.get("entry_requirements"),
                 row.get("method_of_assessment"), row.get("course_overview"),
                 row.get("category_id"), level, ctype,
                 row.get("qualification_spec_current"), now, existing["id"]),
            )
            return "updated"
        c.execute(
            """INSERT INTO courses (course_name, course_url, spec_url, entry_requirements,
               method_of_assessment, course_overview, category_id, level, course_type,
               qualification_spec_current, updated_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
            (row["course_name"], row.get("course_url"), row.get("spec_url"),
             row.get("entry_requirements"), row.get("method_of_assessment"),
             row.get("course_overview"), row.get("category_id"), level, ctype,
             row.get("qualification_spec_current"), now),
        )
        return "inserted"


def all_courses() -> list:
    with get_conn() as c:
        return [dict(r) for r in c.execute("SELECT * FROM courses ORDER BY course_name")]


def get_course(course_id: int) -> dict:
    with get_conn() as c:
        r = c.execute("SELECT * FROM courses WHERE id=?", (course_id,)).fetchone()
        return dict(r) if r else {}


def delete_course(course_id: int):
    with get_conn() as c:
        c.execute("DELETE FROM reports WHERE course_id=?", (course_id,))
        c.execute("DELETE FROM validation_reports WHERE course_id=?", (course_id,))
        c.execute("DELETE FROM courses WHERE id=?", (course_id,))


def update_course_fields(course_id: int, fields: dict):
    if not fields:
        return
    sets = ", ".join(f"{k}=?" for k in fields)
    vals = list(fields.values()) + [datetime.now().isoformat(timespec="seconds"), course_id]
    with get_conn() as c:
        c.execute(f"UPDATE courses SET {sets}, updated_at=? WHERE id=?", vals)


def filter_options() -> dict:
    """Distinct Category ID / Level / Type values — generated dynamically
    from the imported tracker sheet (the courses table)."""
    with get_conn() as c:
        cats = [r[0] for r in c.execute(
            "SELECT DISTINCT category_id FROM courses WHERE category_id IS NOT NULL "
            "AND category_id != '' ORDER BY category_id")]
        levels = [r[0] for r in c.execute(
            "SELECT DISTINCT level FROM courses WHERE level IS NOT NULL "
            "AND level != '' ORDER BY CAST(level AS INTEGER)")]
        types = [r[0] for r in c.execute(
            "SELECT DISTINCT course_type FROM courses WHERE course_type IS NOT NULL "
            "AND course_type != '' ORDER BY course_type")]
    return {"category_ids": cats, "levels": levels, "types": types}


def courses_filtered(category_id=None, level=None, course_type=None) -> list:
    q, args = "SELECT * FROM courses WHERE 1=1", []
    if category_id and category_id != "All":
        q += " AND category_id = ?"; args.append(category_id)
    if level and level != "All":
        q += " AND level = ?"; args.append(level)
    if course_type and course_type != "All":
        q += " AND course_type = ?"; args.append(course_type)
    q += " ORDER BY course_name"
    with get_conn() as c:
        return [dict(r) for r in c.execute(q, args)]


def save_spec_text(course_id: int, text: str):
    with get_conn() as c:
        c.execute("UPDATE courses SET spec_text=?, updated_at=? WHERE id=?",
                  (text, datetime.now().isoformat(timespec="seconds"), course_id))


def save_spec_fields(course_id: int, filename=None, spec_text=None,
                     entry=None, qual=None, assess=None):
    fields = {}
    if filename is not None: fields["spec_filename"] = filename
    if spec_text is not None: fields["spec_text"] = spec_text
    if entry is not None: fields["spec_entry_requirements"] = entry
    if qual is not None: fields["spec_qualification"] = qual
    if assess is not None: fields["spec_assessment"] = assess
    update_course_fields(course_id, fields)


# ── specification documents (one row per UNIQUE document) ───────────

def sync_spec_documents() -> dict:
    """Create a spec_documents row for every distinct specification URL in the
    courses table and link each course to its document (courses.spec_doc_id).
    Detects when multiple courses share the same document so it is only
    processed once. Returns counts."""
    now = datetime.now().isoformat(timespec="seconds")
    created = 0
    with get_conn() as c:
        urls = [r[0] for r in c.execute(
            "SELECT DISTINCT spec_url FROM courses "
            "WHERE spec_url IS NOT NULL AND TRIM(spec_url) != ''")]
        for url in urls:
            u = url.strip()
            row = c.execute("SELECT id FROM spec_documents WHERE spec_url=?", (u,)).fetchone()
            if not row:
                c.execute("INSERT INTO spec_documents (spec_url, updated_at) VALUES (?,?)",
                          (u, now))
                created += 1
        # link every course to its document
        c.execute("""UPDATE courses SET spec_doc_id =
                       (SELECT d.id FROM spec_documents d
                        WHERE d.spec_url = TRIM(courses.spec_url))
                     WHERE spec_url IS NOT NULL AND TRIM(spec_url) != ''""")
        total = c.execute("SELECT COUNT(*) FROM spec_documents").fetchone()[0]
        processed = c.execute(
            "SELECT COUNT(*) FROM spec_documents WHERE status='processed'").fetchone()[0]
    return {"created": created, "total": total, "processed": processed}


def all_spec_docs() -> list:
    """Every specification document with the number of courses that use it."""
    with get_conn() as c:
        return [dict(r) for r in c.execute(
            """SELECT d.*, COUNT(cs.id) AS course_count
               FROM spec_documents d
               LEFT JOIN courses cs ON cs.spec_doc_id = d.id
               GROUP BY d.id ORDER BY d.spec_url""")]


def get_spec_doc(doc_id) -> dict:
    if not doc_id:
        return {}
    with get_conn() as c:
        r = c.execute("SELECT * FROM spec_documents WHERE id=?", (doc_id,)).fetchone()
        return dict(r) if r else {}


def spec_doc_data(doc: dict) -> dict:
    """Parse the stored structured JSON of a processed document."""
    try:
        return json.loads(doc.get("extracted_json") or "{}")
    except Exception:
        return {}


def save_spec_doc(doc_id: int, fields: dict):
    if not fields:
        return
    sets = ", ".join(f"{k}=?" for k in fields)
    vals = list(fields.values()) + [datetime.now().isoformat(timespec="seconds"), doc_id]
    with get_conn() as c:
        c.execute(f"UPDATE spec_documents SET {sets}, updated_at=? WHERE id=?", vals)


def courses_for_doc(doc_id: int) -> list:
    with get_conn() as c:
        return [dict(r) for r in c.execute(
            "SELECT * FROM courses WHERE spec_doc_id=? ORDER BY course_name", (doc_id,))]


def propagate_doc_to_courses(doc_id: int):
    """Mirror the document's extracted key sections onto every linked course's
    legacy spec_* columns (keeps older views & the offline fallback working)."""
    doc = get_spec_doc(doc_id)
    data = spec_doc_data(doc)
    if not data:
        return
    qual = data.get("qualification_specification_requirements") or ""
    head = " · ".join(x for x in [data.get("qualification_name"),
                                  f"Level {data.get('qualification_level')}" if data.get("qualification_level") else "",
                                  data.get("qualification_type")] if x)
    qual_full = (head + "\n" + qual).strip() if head else qual
    with get_conn() as c:
        c.execute("""UPDATE courses SET spec_filename=?, spec_entry_requirements=?,
                     spec_qualification=?, spec_assessment=?, updated_at=?
                     WHERE spec_doc_id=?""",
                  (doc.get("filename") or doc.get("spec_url"),
                   data.get("entry_requirements") or "",
                   qual_full,
                   data.get("method_of_assessment") or "",
                   datetime.now().isoformat(timespec="seconds"), doc_id))


def save_report(course_id, fields, status, errors, summary, rewrites=None):
    with get_conn() as c:
        c.execute(
            """INSERT INTO reports (course_id, checked_at, fields_checked, status,
               errors_json, rewrites_json, summary)
               VALUES (?,?,?,?,?,?,?)""",
            (course_id, datetime.now().isoformat(timespec="seconds"),
             json.dumps(fields), status, json.dumps(errors, ensure_ascii=False),
             json.dumps(rewrites or [], ensure_ascii=False), summary),
        )


def latest_reports() -> list:
    """Latest report per course, joined with the course name."""
    q = """
        SELECT r.*, c.course_name, c.course_url
        FROM reports r
        JOIN courses c ON c.id = r.course_id
        WHERE r.id IN (SELECT MAX(id) FROM reports GROUP BY course_id)
        ORDER BY c.course_name
    """
    with get_conn() as c:
        return [dict(x) for x in c.execute(q)]


def save_validation_report(course_id: int, checked_by: str, status: str,
                           summary: str, results: dict):
    with get_conn() as c:
        c.execute(
            """INSERT INTO validation_reports
               (course_id, checked_at, checked_by, status, summary, results_json)
               VALUES (?,?,?,?,?,?)""",
            (course_id, datetime.now().isoformat(timespec="seconds"), checked_by,
             status, summary, json.dumps(results, ensure_ascii=False)),
        )


def latest_validation_reports() -> list:
    q = """
        SELECT v.*, c.course_name, c.category_id, c.level, c.course_type
        FROM validation_reports v
        JOIN courses c ON c.id = v.course_id
        WHERE v.id IN (SELECT MAX(id) FROM validation_reports GROUP BY course_id)
        ORDER BY c.course_name
    """
    with get_conn() as c:
        return [dict(x) for x in c.execute(q)]


# ═══════════════════════════════════════════════════════════════════
#  FETCHING & EXTRACTION
# ═══════════════════════════════════════════════════════════════════

def fetch_url(url: str, timeout=30) -> requests.Response:
    return requests.get(url, headers=USER_AGENT, timeout=timeout, allow_redirects=True)


def extract_page_text(url: str, max_chars=12000) -> str:
    """Fetch a live course page and return its readable text."""
    resp = fetch_url(url)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "form"]):
        tag.decompose()
    main = soup.find("main") or soup.find("article") or soup.body or soup
    text = re.sub(r"\n{3,}", "\n\n", main.get_text("\n", strip=True))
    return text[:max_chars]


def _direct_download_url(url: str) -> str:
    """Convert share links (e.g. Google Drive /view links) into direct-download
    URLs so the actual document is fetched instead of a viewer page."""
    m = re.search(r"drive\.google\.com/file/d/([^/?#]+)", url or "")
    if m:
        return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    m = re.search(r"drive\.google\.com/open\?id=([^&#]+)", url or "")
    if m:
        return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    return url


def extract_spec_text(url: str, max_chars=15000) -> str:
    """Extract text from a specification document — PDF, Word or web page."""
    resp = fetch_url(_direct_download_url(url), timeout=60)
    resp.raise_for_status()
    ctype = resp.headers.get("Content-Type", "").lower()
    head = resp.content[:8]
    if head.startswith(b"%PDF") or "pdf" in ctype or url.lower().split("?")[0].endswith(".pdf"):
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(resp.content)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
                if sum(len(p) for p in parts) > max_chars:
                    break
        return "\n".join(parts)[:max_chars]
    if head.startswith(b"PK") and ("officedocument" in ctype
                                   or url.lower().split("?")[0].endswith(".docx")
                                   or b"word/" in resp.content[:4000]):
        d = Document(io.BytesIO(resp.content))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(p for p in parts if p and p.strip())[:max_chars]
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    return soup.get_text("\n", strip=True)[:max_chars]


def read_uploaded_spec(file) -> str:
    """Read an uploaded specification document (.pdf or .docx) into plain text."""
    name = file.name.lower()
    data = file.read()
    if name.endswith(".pdf"):
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    if name.endswith(".docx"):
        d = Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(p for p in parts if p and p.strip())
    return data.decode("utf-8", errors="ignore")


# Headings recognised inside qualification specification documents.
SECTION_PATTERNS = {
    "entry": r"entry\s+requirements?|entry\s+criteria|admission\s+requirements?",
    "assessment": r"method(?:s)?\s+of\s+assessment|assessment\s+method(?:s|ology)?"
                  r"|\bassessment\b(?:\s+and\s+grading)?",
    "qualification": r"qualification\s+(?:specification|overview|summary|objective|purpose"
                     r"|structure|details?)|about\s+(?:this|the)\s+qualification",
}

_HEADING_STOP = (
    r"(?:^|\n)\s*(?:\d+(?:\.\d+)*\s+)?"
    r"(entry\s+requirements?|entry\s+criteria|admission\s+requirements?"
    r"|method(?:s)?\s+of\s+assessment|assessment\s+method(?:s|ology)?"
    r"|qualification\s+(?:specification|overview|summary|objective|purpose|structure|details?)"
    r"|about\s+(?:this|the)\s+qualification"
    r"|progression|grading|units?|guided\s+learning|total\s+qualification\s+time"
    r"|course\s+content|learning\s+outcomes?|funding|centre|appendix|introduction"
    r"|support|resources|contact)\b"
)


def extract_spec_sections(text: str, max_chars: int = 2500) -> dict:
    """Heuristically pull Entry Requirements / Method of Assessment /
    Qualification Specification sections out of a spec document's text."""
    out = {"entry": "", "assessment": "", "qualification": ""}
    if not text:
        return out
    for key, pattern in SECTION_PATTERNS.items():
        m = re.search(rf"(?:^|\n)\s*(?:\d+(?:\.\d+)*\s+)?(?:{pattern})\s*:?\s*\n?",
                      text, re.I)
        if not m:
            continue
        start = m.end()
        stop = re.search(_HEADING_STOP, text[start:], re.I)
        chunk = text[start:start + stop.start()] if stop else text[start:]
        chunk = re.sub(r"\n{3,}", "\n\n", chunk).strip()
        out[key] = chunk[:max_chars]
    # fallback for the qualification summary: use the document opening
    if not out["qualification"]:
        head = text.strip()[:1200]
        out["qualification"] = re.sub(r"\n{3,}", "\n\n", head)
    return out


# ═══════════════════════════════════════════════════════════════════
#  OPENROUTER
# ═══════════════════════════════════════════════════════════════════

def call_openrouter(prompt: str, system: str, api_key: str, model: str, temperature=0.0) -> str:
    resp = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:8501",
            "X-Title": "SLC Course Content Checker",
        },
        json={
            "model": model,
            "temperature": temperature,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": prompt},
            ],
        },
        timeout=180,
    )
    resp.raise_for_status()
    data = resp.json()
    return data["choices"][0]["message"]["content"]


def parse_json_reply(raw: str) -> dict:
    """Robustly pull a JSON object out of an LLM reply."""
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.S)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", raw, flags=re.S)
        if m:
            return json.loads(m.group(0))
        raise


AI_SECTION_SYSTEM = (
    "You extract sections from UK qualification specification documents. "
    "You reply ONLY with valid JSON — no markdown, no commentary."
)

AI_SECTION_PROMPT = """From the qualification specification text below, extract these three sections verbatim (clean up broken line-wrapping but do not reword):

1. "entry_requirements" — the entry requirements / entry criteria
2. "qualification_specification" — the qualification overview / specification summary (title, level, size, purpose)
3. "method_of_assessment" — how the qualification is assessed

If a section is genuinely absent, use an empty string.

Reply with EXACTLY this JSON:
{{"entry_requirements": "...", "qualification_specification": "...", "method_of_assessment": "..."}}

=== SPECIFICATION TEXT ===
{text}
"""


def ai_extract_sections(text: str, api_key: str, model: str) -> dict:
    raw = call_openrouter(AI_SECTION_PROMPT.format(text=text[:14000]),
                          AI_SECTION_SYSTEM, api_key, model)
    data = parse_json_reply(raw)
    return {
        "entry": str(data.get("entry_requirements", "")).strip(),
        "qualification": str(data.get("qualification_specification", "")).strip(),
        "assessment": str(data.get("method_of_assessment", "")).strip(),
    }


# ═══════════════════════════════════════════════════════════════════
#  ONE-TIME SPECIFICATION DOCUMENT PROCESSING
#  Each unique specification document is read + AI-extracted ONCE, the
#  structured result is stored as JSON in the database and REUSED for
#  every course and every Run Check — no re-extraction at check time.
# ═══════════════════════════════════════════════════════════════════

# field key → human label (order used across the UI)
SPECDOC_FIELDS = [
    ("qualification_name",   "Qualification Name"),
    ("qualification_level",  "Qualification Level"),
    ("qualification_type",   "Qualification Type"),
    ("entry_requirements",   "Entry Requirements"),
    ("method_of_assessment", "Method of Assessment"),
    ("qualification_specification_requirements", "Qualification Specification Requirements"),
    ("learning_outcomes",    "Learning Outcomes"),
    ("mandatory_units",      "Mandatory Units"),
    ("other_information",    "Other Relevant Information"),
]

AI_SPECDOC_SYSTEM = (
    "You extract structured data from UK qualification specification documents. "
    "You reply ONLY with valid JSON — no markdown, no commentary."
)

AI_SPECDOC_PROMPT = """Extract ALL of the following from the qualification specification text below. Copy wording faithfully (fix broken line-wrapping but do not reword). Use "" for absent strings and [] for absent lists.

Reply with EXACTLY this JSON shape:
{{
  "qualification_name": "the official qualification title",
  "qualification_level": "e.g. 3",
  "qualification_type": "Award | Certificate | Diploma",
  "entry_requirements": "the full entry requirements / entry criteria section",
  "method_of_assessment": "the full method of assessment section",
  "qualification_specification_requirements": "the qualification overview / specification requirements: purpose, size (TQT / GLH / credits), structure, rules of combination and any key requirements",
  "learning_outcomes": ["each learning outcome as one string"],
  "mandatory_units": ["each mandatory unit as one string, e.g. 'Unit 1: ... (code, credits)'"],
  "other_information": "anything else useful for validating a course page against this specification (grading, progression, age ranges, etc.)"
}}

=== SPECIFICATION TEXT ===
{text}
"""


def ai_extract_spec_json(text: str, api_key: str, model: str) -> dict:
    raw = call_openrouter(AI_SPECDOC_PROMPT.format(text=text[:24000]),
                          AI_SPECDOC_SYSTEM, api_key, model)
    data = parse_json_reply(raw)
    out = {}
    for key, _ in SPECDOC_FIELDS:
        v = data.get(key, [] if key in ("learning_outcomes", "mandatory_units") else "")
        if key in ("learning_outcomes", "mandatory_units"):
            out[key] = [str(x).strip() for x in (v or []) if str(x).strip()]
        else:
            out[key] = str(v or "").strip()
    return out


def heuristic_spec_json(text: str) -> dict:
    """Offline fallback — builds the structured record from heading heuristics."""
    secs = extract_spec_sections(text, max_chars=4000)
    head = (text or "").strip().splitlines()
    title = next((l.strip() for l in head if len(l.strip()) > 12), "")[:200]
    m_lvl = re.search(r"\blevel\s*(\d)\b", text or "", re.I)
    m_typ = re.search(r"\b(award|certificate|diploma)\b", text or "", re.I)
    units = re.findall(r"(?im)^\s*(unit\s+\d+[^\n]{0,140})$", text or "")
    los = re.findall(r"(?im)^\s*((?:LO\s*\d+|learning outcome\s*\d+)[^\n]{0,160})$", text or "")
    return {
        "qualification_name": title,
        "qualification_level": m_lvl.group(1) if m_lvl else "",
        "qualification_type": m_typ.group(1).title() if m_typ else "",
        "entry_requirements": secs["entry"],
        "method_of_assessment": secs["assessment"],
        "qualification_specification_requirements": secs["qualification"],
        "learning_outcomes": [l.strip() for l in los][:40],
        "mandatory_units": [u.strip() for u in units][:40],
        "other_information": "",
    }


def process_spec_document(doc_id: int, api_key: str, model: str,
                          force: bool = False,
                          uploaded_text: str = None, uploaded_name: str = None) -> dict:
    """Read + AI-extract ONE specification document and store the result.

    - Skips work entirely when the document is already processed (reuse),
      unless `force` is set or replacement content is supplied.
    - Skips re-extraction when a re-fetched document is unchanged (hash check).
    Returns {"status": "processed"|"reused"|"unchanged"|"error", "detail": str}.
    """
    doc = get_spec_doc(doc_id)
    if not doc:
        return {"status": "error", "detail": "Document not found."}

    if doc.get("status") == "processed" and not force and uploaded_text is None:
        return {"status": "reused", "detail": "Already processed — stored data reused."}

    # 1) obtain the document text (upload beats URL fetch)
    try:
        if uploaded_text is not None:
            text, fname = uploaded_text, (uploaded_name or doc.get("filename"))
        else:
            if not doc.get("spec_url"):
                raise ValueError("No specification URL recorded and no file uploaded.")
            text = extract_spec_text(doc["spec_url"], max_chars=30000)
            fname = doc["spec_url"]
        if not (text or "").strip():
            raise ValueError("The document contained no extractable text.")
    except Exception as e:
        save_spec_doc(doc_id, {"status": "error", "error": str(e)})
        return {"status": "error", "detail": f"Could not read the document: {e}"}

    # 2) unchanged? → keep the existing extraction (unless forced)
    h = hashlib.sha256(text.encode("utf-8", "ignore")).hexdigest()
    if (doc.get("status") == "processed" and doc.get("content_hash") == h
            and not force and doc.get("extracted_json")):
        save_spec_doc(doc_id, {"filename": fname, "error": None})
        return {"status": "unchanged", "detail": "Document unchanged — existing extraction kept."}

    # 3) extract structured data (AI, with heuristic fallback)
    method = "heuristic"
    data = None
    if api_key:
        try:
            data = ai_extract_spec_json(text, api_key, model)
            method = "ai"
        except Exception:
            data = None
    if data is None:
        data = heuristic_spec_json(text)

    now = datetime.now().isoformat(timespec="seconds")
    save_spec_doc(doc_id, {
        "filename": fname, "raw_text": text[:60000],
        "extracted_json": json.dumps(data, ensure_ascii=False),
        "status": "processed", "method": method, "error": None,
        "content_hash": h, "processed_at": now,
    })
    propagate_doc_to_courses(doc_id)
    return {"status": "processed", "detail": f"Processed with {method} extraction."}


def spec_data_for_prompt(data: dict, limit: int = 12000) -> str:
    """Compact plain-text rendering of the stored structured data for the
    validation prompt — keeps AI token usage low (no raw document text)."""
    parts = []
    for key, label in SPECDOC_FIELDS:
        v = data.get(key)
        if not v:
            continue
        if isinstance(v, list):
            v = "\n".join(f"- {x}" for x in v)
        parts.append(f"## {label}\n{v}")
    return "\n\n".join(parts)[:limit]


# ═══════════════════════════════════════════════════════════════════
#  COURSE CHECKING  (existing "Other Checks" — unchanged functionality)
# ═══════════════════════════════════════════════════════════════════

CHECK_SYSTEM = (
    "You are a meticulous content auditor AND professional copywriter for South London College. "
    "You compare the LIVE course web page against the internal TRACKER content and the OFFICIAL "
    "qualification SPECIFICATION. You flag factual mismatches, missing information, outdated "
    "details and misleading statements. You also write replacement wording that is 100% original: "
    "factually faithful to the specification but never copying its sentences or distinctive "
    "phrases. You reply ONLY with valid JSON — no markdown, no commentary."
)

CHECK_PROMPT = """Compare the three sources below for the course "{name}" and audit ONLY these fields: {fields}.

For each audited field decide whether the LIVE PAGE content is accurate and consistent with the TRACKER and the SPECIFICATION. Ignore styling/wording differences that do not change meaning. Flag:
- factual mismatches (grades, units, hours, assessment methods, requirements)
- content present in tracker/spec but missing on the live page
- content on the live page contradicted by the specification

Reply with EXACTLY this JSON shape:
{{
  "status": "pass" or "errors",
  "summary": "one-sentence overall verdict",
  "errors": [
    {{
      "field": "Entry Requirements | Method of Assessment | Course Overview",
      "severity": "high | medium | low",
      "issue": "clear description of the problem",
      "live_content": "the problematic text on the live page (short quote or 'missing')",
      "expected_content": "what tracker/spec says it should be",
      "suggested_fix": "the exact corrected text or action to take"
    }}
  ],
  "rewrites": [
    {{
      "field": "Entry Requirements | Method of Assessment | Course Overview",
      "suggested_wording": "a complete, publish-ready rewrite of this field for the website"
    }}
  ]
}}

RULES FOR "rewrites" (very important):
- Provide one rewrite for EVERY field that has at least one error (empty list if none).
- The rewrite must be publish-ready website copy in UK English: complete, clear, professional.
- It must be factually accurate to the SPECIFICATION and TRACKER (grades, units, hours, methods).
- ZERO PLAGIARISM: write in completely original words. Do NOT reuse sentences, clauses, or
  distinctive phrases of 4+ consecutive words from the specification, the live page, or the
  tracker. Restructure sentences and use your own vocabulary while keeping every fact identical.
- Do not invent facts that are not supported by the specification or tracker.

If everything is accurate return status "pass" and empty errors and rewrites lists.

=== TRACKER CONTENT ===
{tracker}

=== LIVE PAGE CONTENT ===
{live}

=== OFFICIAL SPECIFICATION EXTRACT ===
{spec}
"""


def _shingles(text: str, n: int = 4) -> set:
    words = re.findall(r"[a-z0-9']+", (text or "").lower())
    return {tuple(words[i:i + n]) for i in range(len(words) - n + 1)}


def originality_score(candidate: str, sources: list, n: int = 4) -> float:
    """% of the candidate's 4-word phrases that do NOT appear in any source.
    100 = fully original wording, 0 = copied verbatim."""
    cand = _shingles(candidate, n)
    if not cand:
        return 100.0
    src = set()
    for s in sources:
        src |= _shingles(s, n)
    overlap = len(cand & src) / len(cand)
    return round((1 - overlap) * 100, 1)


REWRITE_RETRY_PROMPT = """Rewrite the text below so that it keeps EXACTLY the same facts but shares NO phrase of 4 or more consecutive words with any of the reference texts. Use different sentence structures and vocabulary. UK English, publish-ready website copy. Reply with ONLY the rewritten text — no JSON, no quotes, no commentary.

=== TEXT TO REWRITE ===
{text}

=== REFERENCE TEXTS IT MUST NOT COPY FROM ===
{refs}
"""


def ensure_original(wording: str, sources: list, api_key: str, model: str,
                    threshold: float = 90.0, max_retries: int = 2) -> tuple:
    """Verify a rewrite is plagiarism-free; if not, ask the model to re-word it.
    Returns (final_wording, originality_score)."""
    score = originality_score(wording, sources)
    tries = 0
    while score < threshold and tries < max_retries:
        try:
            refs = "\n\n---\n\n".join(s[:4000] for s in sources if s)
            wording = call_openrouter(
                REWRITE_RETRY_PROMPT.format(text=wording, refs=refs),
                "You are a professional copywriter. You paraphrase with zero plagiarism.",
                api_key, model, temperature=0.7,
            ).strip()
            score = originality_score(wording, sources)
        except Exception:
            break
        tries += 1
    return wording, score


def check_course(course: dict, fields: list, api_key: str, model: str) -> dict:
    """Run a full 3-way check for one course. Returns a report dict."""
    report = {"course_id": course["id"], "course_name": course["course_name"],
              "status": "failed", "summary": "", "errors": [], "rewrites": []}

    # 1) live page
    try:
        live = extract_page_text(course["course_url"]) if course.get("course_url") else ""
        if not live:
            raise ValueError("No course URL / empty page")
    except Exception as e:
        report["summary"] = f"Could not fetch live page: {e}"
        return report

    # 2) spec text (extract on the fly if missing)
    spec = course.get("spec_text") or ""
    if not spec and course.get("spec_url"):
        try:
            spec = extract_spec_text(course["spec_url"])
            save_spec_text(course["id"], spec)
        except Exception:
            spec = "(specification unavailable)"
    spec = spec or "(no specification provided)"

    # 3) tracker content for the selected fields
    tracker_parts = []
    for key in fields:
        label = FIELD_OPTIONS[key]
        tracker_parts.append(f"[{label}]\n{course.get(key) or '(not provided in tracker)'}")
    tracker = "\n\n".join(tracker_parts)

    prompt = CHECK_PROMPT.format(
        name=course["course_name"],
        fields=", ".join(FIELD_OPTIONS[k] for k in fields),
        tracker=tracker, live=live, spec=spec,
    )

    try:
        raw = call_openrouter(prompt, CHECK_SYSTEM, api_key, model)
        data = parse_json_reply(raw)
        report["status"] = "errors" if data.get("errors") else "pass"
        if data.get("status") == "pass":
            report["status"] = "pass"
        report["summary"] = data.get("summary", "")
        report["errors"] = data.get("errors", [])

        # ── wording suggestions with zero-plagiarism verification ──
        sources = [spec, live, tracker]
        rewrites = []
        for rw in data.get("rewrites", []) or []:
            wording = str(rw.get("suggested_wording", "")).strip()
            if not wording:
                continue
            wording, score = ensure_original(wording, sources, api_key, model)
            rewrites.append({
                "field": rw.get("field", ""),
                "suggested_wording": wording,
                "originality": score,
            })
        report["rewrites"] = rewrites
    except Exception as e:
        report["summary"] = f"AI check failed: {e}"
    return report


# ═══════════════════════════════════════════════════════════════════
#  NEW — SPEC VALIDATION (Run Check → Content Check)
# ═══════════════════════════════════════════════════════════════════

VALIDATE_SYSTEM = (
    "You are a meticulous quality auditor for South London College. You compare a live "
    "COURSE PAGE (plus the internal course record) against the OFFICIAL QUALIFICATION "
    "SPECIFICATION data and flag incorrect wording, incorrect information, missing "
    "information, mismatched requirements and grammar issues. You reply ONLY with "
    "valid JSON — no markdown, no commentary."
)

VALIDATE_PROMPT = """Validate the course "{name}" (Level {level} {ctype}, ref {number}).

=== OFFICIAL QUALIFICATION SPECIFICATION (extracted, authoritative) ===
{spec}

=== LIVE COURSE PAGE CONTENT ===
{page}

=== INTERNAL COURSE RECORD (tracker) ===
Entry Requirements: {rec_entry}
Method of Assessment: {rec_assess}

Compare the course content against the specification and report on EXACTLY these sections:

1. "Qualification Specification" — the qualification title, level, type, size (TQT/GLH/credits), structure and key specification requirements stated on the course page / record must match the specification. Also verify learning outcomes and mandatory units mentioned by the course are correct.
2. "Entry Requirement" — every requirement in the specification must be present and correct; flag anything missing, contradictory or extra.
3. "Method of Assessment" — compared on WORDING ONLY: the course wording must faithfully match the specification's assessment wording (paraphrase is fine, but nothing incorrect, extra or missing).
4. "Content & Grammar" — any incorrect information or notable grammar/wording problems in the course content relevant to the specification (only clear issues; do not nitpick style).

Rules:
- Prefer the LIVE COURSE PAGE as the "current" text for each section; fall back to the internal record when the page lacks that section.
- In each section, set "current" to the exact course text you assessed (short excerpt, max ~120 words).
- If the course text for a section is empty → one issue of type "missing".
- If the specification lacks the data needed to verify a section → one issue of type "mismatch" saying it could not be verified.
- Number issues per section starting at 1. Keep each error and recommendation to one or two clear sentences. Recommendations must state the corrected wording where possible.
- issue "type" must be one of: incorrect_wording | incorrect_information | missing | mismatch | grammar.

Reply with EXACTLY this JSON shape:
{{
  "status": "pass" or "errors",
  "summary": "one or two sentence overall verdict",
  "sections": [
    {{
      "section": "Qualification Specification | Entry Requirement | Method of Assessment | Content & Grammar",
      "status": "correct | incorrect_wording | missing | mismatch",
      "current": "the course text assessed for this section",
      "issues": [
        {{
          "type": "incorrect_wording | incorrect_information | missing | mismatch | grammar",
          "error": "clear description of the error identified",
          "recommendation": "the recommended action / corrected wording"
        }}
      ]
    }}
  ]
}}

If every section is fully correct, return status "pass" with all sections "correct" and empty issue lists.
"""


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _fallback_compare(label: str, current: str, spec: str) -> dict:
    """Deterministic comparison used when no API key is configured."""
    section = {"section": label, "status": "correct", "issues": []}
    cur, sp = _norm(current), _norm(spec)
    if not sp:
        section["status"] = "mismatch"
        section["issues"].append({
            "type": "mismatch",
            "error": f"The specification document does not contain an extractable "
                     f"'{label}' section, so this section could not be verified.",
            "recommendation": "Upload or re-extract the specification document and make "
                              "sure the section is captured in 📑 Spec Documents.",
        })
        return section
    if not cur:
        section["status"] = "missing"
        section["issues"].append({
            "type": "missing",
            "error": f"The course record has no '{label}' content, but the "
                     f"specification defines it.",
            "recommendation": f"Add the following to the course record: "
                              f"{spec.strip()[:400]}",
        })
        return section
    sim = SequenceMatcher(None, cur, sp).ratio()
    # sentences in the spec that are absent from the current text
    missing_bits = []
    for sent in re.split(r"(?<=[.;•])\s+|\n+", spec):
        s = _norm(sent)
        if len(s) < 25:
            continue
        if SequenceMatcher(None, s, cur).find_longest_match(
                0, len(s), 0, len(cur)).size < min(30, int(len(s) * .55)) and s not in cur:
            missing_bits.append(sent.strip())
    if sim >= 0.92 and not missing_bits:
        return section
    n = 0
    if missing_bits:
        n += 1
        section["status"] = "mismatch"
        section["issues"].append({
            "type": "missing",
            "error": "Information stated in the specification appears to be missing "
                     "or worded differently in the current text: "
                     + " | ".join(missing_bits[:3])[:400],
            "recommendation": "Update the current text so it includes/matches the "
                              "specification wording for the points above.",
        })
    if sim < 0.6:
        section["status"] = "incorrect_wording"
        section["issues"].append({
            "type": "incorrect_wording",
            "error": f"The current wording differs substantially from the "
                     f"specification (similarity {sim:.0%}).",
            "recommendation": f"Align the wording with the specification: "
                              f"{spec.strip()[:400]}",
        })
    if not section["issues"]:
        section["issues"].append({
            "type": "incorrect_wording",
            "error": f"The current wording deviates from the specification "
                     f"(similarity {sim:.0%}).",
            "recommendation": "Review the section against the specification and align "
                              "the wording.",
        })
        section["status"] = "incorrect_wording"
    return section


def validate_course_vs_spec(course: dict, spec_data: dict, page_text: str,
                            api_key: str, model: str) -> dict:
    """Compare the course content (live page + tracker record) against the
    STORED specification data — the document is never re-read here.
    Returns {status, summary, sections:[{section,status,current,issues:[...]}]}."""
    result = {"status": "pass", "summary": "", "sections": []}

    if api_key and spec_data:
        prompt = VALIDATE_PROMPT.format(
            name=course.get("course_name", ""),
            level=course.get("level") or "?",
            ctype=course.get("course_type") or "?",
            number=course.get("category_id") or "—",
            spec=spec_data_for_prompt(spec_data),
            page=(page_text or "(course page unavailable)")[:9000],
            rec_entry=(course.get("entry_requirements") or "(empty)")[:1500],
            rec_assess=(course.get("method_of_assessment") or "(empty)")[:1500],
        )
        try:
            data = parse_json_reply(call_openrouter(prompt, VALIDATE_SYSTEM, api_key, model))
            for s in data.get("sections", []):
                result["sections"].append({
                    "section": str(s.get("section", "Section")),
                    "current": str(s.get("current", "") or ""),
                    "status": s.get("status", "correct"),
                    "issues": s.get("issues", []) or [],
                })
            result["summary"] = data.get("summary", "")
            result["status"] = ("errors" if any(sec["issues"] for sec in result["sections"])
                                else "pass")
            if result["sections"]:
                return result
        except Exception as e:
            result["summary"] = f"AI validation unavailable ({e}) — used built-in comparison. "
            result["sections"] = []

    # deterministic fallback (no API key / AI failure) — compares against the
    # stored structured data, never re-reads the document
    qual_head = " · ".join(x for x in [
        spec_data.get("qualification_name"),
        f"Level {spec_data.get('qualification_level')}" if spec_data.get("qualification_level") else "",
        spec_data.get("qualification_type")] if x)
    spec_qual = (qual_head + "\n" + (spec_data.get("qualification_specification_requirements") or "")).strip()
    src = page_text or ""
    pairs = [
        ("Qualification Specification",
         course.get("qualification_spec_current") or course.get("course_name") or src[:1200],
         spec_qual),
        ("Entry Requirement",
         course.get("entry_requirements") or "",
         spec_data.get("entry_requirements") or ""),
        ("Method of Assessment",
         course.get("method_of_assessment") or "",
         spec_data.get("method_of_assessment") or ""),
    ]
    for label, current, spec in pairs:
        s = _fallback_compare(label, current, spec)
        s.update({"current": current, "spec": spec})
        result["sections"].append(s)
    n_issues = sum(len(s["issues"]) for s in result["sections"])
    result["status"] = "errors" if n_issues else "pass"
    result["summary"] += ("No issues detected — all sections match the specification."
                          if not n_issues else
                          f"{n_issues} issue(s) detected across "
                          f"{sum(1 for s in result['sections'] if s['issues'])} section(s).")
    return result


# ═══════════════════════════════════════════════════════════════════
#  WORD REPORT (errors only — existing export, unchanged)
# ═══════════════════════════════════════════════════════════════════

SEV_COLORS = {"high": RGBColor(0xC6, 0x28, 0x28),
              "medium": RGBColor(0xE6, 0x7E, 0x22),
              "low": RGBColor(0x2E, 0x7D, 0x32)}


def build_word_report(error_reports: list) -> bytes:
    doc = Document()

    title = doc.add_heading("SLC Course Content Error Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Generated {datetime.now().strftime('%d %B %Y, %H:%M')}  ·  "
        f"{len(error_reports)} course(s) with errors"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    for rep in error_reports:
        doc.add_heading(rep["course_name"], level=1)
        if rep.get("course_url"):
            p = doc.add_paragraph()
            r = p.add_run(rep["course_url"])
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        if rep.get("summary"):
            p = doc.add_paragraph(rep["summary"])
            p.runs[0].italic = True

        errors = rep.get("errors") or []
        if not errors:
            doc.add_paragraph("Check failed — see summary above.")
            continue

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Field", "Severity", "Issue", "Live page says", "Suggested solution"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True

        for err in errors:
            cells = table.add_row().cells
            cells[0].text = str(err.get("field", ""))
            sev = str(err.get("severity", "medium")).lower()
            cells[1].text = sev.upper()
            for run in cells[1].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = SEV_COLORS.get(sev, SEV_COLORS["medium"])
            cells[2].text = str(err.get("issue", ""))
            cells[3].text = str(err.get("live_content", ""))
            fix_cell = cells[4]
            fix_cell.text = str(err.get("suggested_fix", ""))
            if err.get("expected_content"):
                p = fix_cell.add_paragraph()
                r = p.add_run(f"Expected: {err['expected_content']}")
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        rewrites = rep.get("rewrites") or []
        if rewrites:
            doc.add_heading("Suggested replacement wording (plagiarism-checked)", level=2)
            for rw in rewrites:
                p = doc.add_paragraph()
                r = p.add_run(f"{rw.get('field','')}  ·  originality {rw.get('originality','—')}%")
                r.bold = True
                r.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)
                doc.add_paragraph(str(rw.get("suggested_wording", "")))

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
#  NEW — VALIDATION REPORT EXPORTS (Word + PDF, screenshot layout)
# ═══════════════════════════════════════════════════════════════════

_DOCX_RED = RGBColor(0xD7, 0x19, 0x20)
_DOCX_GREEN = RGBColor(0x1E, 0x9E, 0x3E)


def _docx_run(p, text, color=None, bold=False, size=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return r


def build_validation_docx(course: dict, result: dict) -> bytes:
    """Word version of the validation report — same layout as the agreed screenshot."""
    doc = Document()

    title = doc.add_heading("Course Validation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"{course.get('course_name','')}  ·  generated "
                            f"{datetime.now().strftime('%d %B %Y, %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ── Level · Type · Number header (red, like the screenshot) ──
    head = doc.add_table(rows=2, cols=3)
    head.style = "Table Grid"
    labels = ["Level", "Type", "Number"]
    values = [str(course.get("level") or "—"),
              str(course.get("course_type") or "—"),
              str(course.get("category_id") or "—")]
    for i in range(3):
        p = head.rows[0].cells[i].paragraphs[0]
        _docx_run(p, labels[i], color=_DOCX_RED, bold=True, size=13)
        head.rows[1].cells[i].text = values[i]
    doc.add_paragraph()

    if result.get("status") == "pass":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_run(p, "✔  NO ERRORS — this course passed all validation checks",
                  color=_DOCX_GREEN, bold=True, size=14)
        doc.add_paragraph()

    for sec in result.get("sections", []):
        p = doc.add_paragraph()
        _docx_run(p, f"Current {sec['section']}", color=_DOCX_RED, bold=True, size=12)

        # boxed current content
        box = doc.add_table(rows=1, cols=1)
        box.style = "Table Grid"
        box.rows[0].cells[0].text = (sec.get("current") or "(no content recorded)")[:1800]

        issues = sec.get("issues") or []
        if not issues:
            p = doc.add_paragraph()
            _docx_run(p, "✔ Correct — matches the specification",
                      color=_DOCX_GREEN, bold=True, size=10)
        for n, iss in enumerate(issues, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _docx_run(p, f"Errors identified {n:02d}:  ", color=_DOCX_RED, bold=True, size=10)
            _docx_run(p, str(iss.get("error", "")), size=10)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(8)
            _docx_run(p2, f"Recommend Action {n:02d}:  ", color=_DOCX_GREEN, bold=True, size=10)
            _docx_run(p2, str(iss.get("recommendation", "")), size=10)
        doc.add_paragraph()

    # ── summary ──
    doc.add_heading("Summary of detected issues", level=1)
    if result.get("summary"):
        p = doc.add_paragraph(result["summary"])
        p.runs[0].italic = True
    total = sum(len(s.get("issues") or []) for s in result.get("sections", []))
    if total == 0:
        p = doc.add_paragraph()
        _docx_run(p, "No issues detected.", color=_DOCX_GREEN, bold=True)
    else:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "Section", "Error identified", "Recommended action"]):
            tbl.rows[0].cells[i].text = h
            for run in tbl.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        k = 0
        for sec in result.get("sections", []):
            for iss in sec.get("issues") or []:
                k += 1
                cells = tbl.add_row().cells
                cells[0].text = str(k)
                cells[1].text = sec["section"]
                cells[2].text = str(iss.get("error", ""))
                cells[3].text = str(iss.get("recommendation", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_validation_pdf(course: dict, result: dict) -> bytes:
    """PDF version of the validation report — same layout as the agreed screenshot."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    red = colors.HexColor(RED)
    green = colors.HexColor(GREEN)
    styles = getSampleStyleSheet()
    s_title = ParagraphStyle("t", parent=styles["Title"], fontSize=18)
    s_sub = ParagraphStyle("sub", parent=styles["Normal"], fontSize=9,
                           textColor=colors.HexColor("#666666"), alignment=TA_CENTER)
    s_red_h = ParagraphStyle("rh", parent=styles["Normal"], fontSize=12,
                             textColor=red, fontName="Helvetica-Bold", spaceAfter=4)
    s_body = ParagraphStyle("b", parent=styles["Normal"], fontSize=9, leading=13)
    s_err = ParagraphStyle("e", parent=styles["Normal"], fontSize=9, leading=12,
                           textColor=red)
    s_act = ParagraphStyle("a", parent=styles["Normal"], fontSize=9, leading=12,
                           textColor=green, spaceAfter=6)
    s_ok = ParagraphStyle("ok", parent=styles["Normal"], fontSize=10,
                          textColor=green, fontName="Helvetica-Bold")
    s_pass = ParagraphStyle("p", parent=styles["Normal"], fontSize=13, alignment=TA_CENTER,
                            textColor=green, fontName="Helvetica-Bold")

    def esc(t):
        return html.escape(str(t or "")).replace("\n", "<br/>")

    story = [Paragraph("Course Validation Report", s_title),
             Paragraph(f"{esc(course.get('course_name'))} · generated "
                       f"{datetime.now().strftime('%d %B %Y, %H:%M')}", s_sub),
             Spacer(1, 8 * mm)]

    # Level / Type / Number header
    head = Table([
        [Paragraph(f'<font color="{RED}"><b>{h}</b></font>', s_body)
         for h in ["Level", "Type", "Number"]],
        [Paragraph(esc(course.get("level") or "—"), s_body),
         Paragraph(esc(course.get("course_type") or "—"), s_body),
         Paragraph(esc(course.get("category_id") or "—"), s_body)],
    ], colWidths=[55 * mm, 55 * mm, 60 * mm])
    head.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#BBBBBB")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FDF2F2")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story += [head, Spacer(1, 6 * mm)]

    if result.get("status") == "pass":
        story += [Paragraph("✔ NO ERRORS — this course passed all validation checks", s_pass),
                  Spacer(1, 6 * mm)]

    for sec in result.get("sections", []):
        story.append(Paragraph(f"Current {esc(sec['section'])}", s_red_h))
        box = Table([[Paragraph(esc((sec.get("current") or "(no content recorded)")[:1800]),
                                s_body)]], colWidths=[170 * mm])
        box.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 1.6, red),
            ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story += [box, Spacer(1, 3 * mm)]

        issues = sec.get("issues") or []
        if not issues:
            story.append(Paragraph("✔ Correct — matches the specification", s_ok))
        for n, iss in enumerate(issues, start=1):
            story.append(Paragraph(f"<b>Errors identified {n:02d}:</b> "
                                   f"{esc(iss.get('error'))}", s_err))
            story.append(Paragraph(f"<b>Recommend Action {n:02d}:</b> "
                                   f"{esc(iss.get('recommendation'))}", s_act))
        story.append(Spacer(1, 6 * mm))

    # summary
    story.append(Paragraph("Summary of detected issues",
                           ParagraphStyle("h1", parent=styles["Heading1"], fontSize=13)))
    if result.get("summary"):
        story.append(Paragraph(esc(result["summary"]),
                               ParagraphStyle("i", parent=s_body, fontName="Helvetica-Oblique")))
        story.append(Spacer(1, 2 * mm))
    rows = [[Paragraph("<b>#</b>", s_body), Paragraph("<b>Section</b>", s_body),
             Paragraph("<b>Error identified</b>", s_body),
             Paragraph("<b>Recommended action</b>", s_body)]]
    k = 0
    for sec in result.get("sections", []):
        for iss in sec.get("issues") or []:
            k += 1
            rows.append([Paragraph(str(k), s_body), Paragraph(esc(sec["section"]), s_body),
                         Paragraph(esc(iss.get("error")), s_body),
                         Paragraph(esc(iss.get("recommendation")), s_body)])
    if k == 0:
        story.append(Paragraph("No issues detected.", s_ok))
    else:
        tbl = Table(rows, colWidths=[8 * mm, 40 * mm, 61 * mm, 61 * mm], repeatRows=1)
        tbl.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6FA")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(tbl)

    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=A4, topMargin=16 * mm, bottomMargin=16 * mm,
                      leftMargin=18 * mm, rightMargin=18 * mm,
                      title="Course Validation Report").build(story)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
#  CONTENT QUALITY REVIEW  (existing "Grammar Check" — unchanged)
# ═══════════════════════════════════════════════════════════════════

QR_SYSTEM = (
    "You are a professional UK-English proofreader and copy editor for a college website. "
    "You review course content for grammar, articles (a/an/the), sentence structure, "
    "capitalisation, proper nouns, spelling, commas and punctuation consistency. "
    "You reply ONLY with valid JSON."
)

QR_PROMPT = """Proofread the text below. Find every issue and classify it into EXACTLY one of these categories:
grammar, article, spelling, punctuation, capitalisation, proper_noun, sentence_structure, consistency

Rules:
- "original" must be an EXACT substring copied verbatim from the text (short — the smallest span that contains the problem).
- "correction" is the fixed version of that span.
- "explanation" is one short sentence.
- Also produce the fully corrected version of the whole text.

Reply with EXACTLY this JSON:
{{
  "issues": [
    {{"category": "...", "original": "...", "correction": "...", "explanation": "..."}}
  ],
  "corrected_text": "..."
}}

=== TEXT TO REVIEW ===
{text}
"""


def run_quality_review(text: str, api_key: str, model: str) -> dict:
    raw = call_openrouter(QR_PROMPT.format(text=text), QR_SYSTEM, api_key, model)
    return parse_json_reply(raw)


def annotate_text_html(text: str, issues: list) -> str:
    """Return HTML with each issue wrapped in a coloured <mark>, numbered like a proofreader's markup."""
    escaped = html.escape(text)
    for n, issue in enumerate(issues, start=1):
        original = html.escape(str(issue.get("original", "")))
        if not original:
            continue
        cat = issue.get("category", "grammar")
        style = QR_CATEGORIES.get(cat, QR_CATEGORIES["grammar"])
        tip = html.escape(f"{style['label']}: {issue.get('correction','')} — {issue.get('explanation','')}")
        mark = (
            f'<mark class="qr-mark" style="background:{style["bg"]};'
            f'border-bottom:2px solid {style["border"]};" title="{tip}">'
            f'<sup class="qr-num" style="background:{style["border"]};">{n}</sup>{original}</mark>'
        )
        escaped = escaped.replace(original, mark, 1)
    return escaped.replace("\n", "<br>")


# ═══════════════════════════════════════════════════════════════════
#  UI — THEME
# ═══════════════════════════════════════════════════════════════════

st.set_page_config(page_title="SLC Course Content Checker", page_icon="🎓",
                   layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,600;9..144,700&family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@500&display=swap');

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

/* ── Hero banner ─────────────────────────────── */
.slc-hero {
  background: linear-gradient(120deg, #10243E 0%, #1D4E89 55%, #2E7D6B 100%);
  border-radius: 18px; padding: 26px 34px; margin-bottom: 6px;
  position: relative; overflow: hidden;
}
.slc-hero:before {
  content:""; position:absolute; right:-40px; top:-40px; width:220px; height:220px;
  background: radial-gradient(circle, rgba(255,229,92,.35), transparent 70%);
}
.slc-hero h1 {
  font-family:'Fraunces', serif; color:#fff; margin:0; font-size:2rem; letter-spacing:.3px;
}
.slc-hero h1 .hl { background:#FFE55C; color:#10243E; padding:0 .35rem; border-radius:6px; }
.slc-hero p { color:#CFE3F5; margin:.45rem 0 0; font-size:.95rem; }

/* ── Stat cards ─────────────────────────────── */
.stat-card {
  border-radius:14px; padding:16px 18px; border:1px solid #E6E9EF;
  background:#fff; box-shadow:0 2px 10px rgba(16,36,62,.06);
}
.stat-card .num { font-family:'Fraunces',serif; font-size:1.9rem; font-weight:700; color:#10243E; line-height:1.1;}
.stat-card .lbl { font-size:.78rem; text-transform:uppercase; letter-spacing:.09em; color:#6B7688; font-weight:600;}
.stat-card.ok   { border-top:4px solid #30A46C; }
.stat-card.err  { border-top:4px solid #E5484D; }
.stat-card.info { border-top:4px solid #1D4E89; }
.stat-card.warn { border-top:4px solid #F5B300; }

/* ── Quality-review markup ───────────────────── */
.qr-paper {
  background:#FFFDF6; border:1px solid #EDE6D2; border-radius:14px;
  padding:26px 30px; line-height:2.05; font-size:1.0rem; color:#2A2F3A;
  box-shadow: 0 3px 14px rgba(16,36,62,.07);
}
.qr-mark { border-radius:4px; padding:1px 3px; cursor:help; position:relative; }
.qr-num {
  color:#fff; font-size:.62rem; font-weight:700; border-radius:999px;
  padding:0 4px; margin-right:2px; position:relative; top:-7px;
}
.qr-legend span {
  display:inline-block; margin:3px 8px 3px 0; padding:3px 10px; border-radius:999px;
  font-size:.78rem; font-weight:600;
}
.issue-card {
  border-radius:12px; border:1px solid #E6E9EF; border-left-width:5px;
  padding:12px 16px; margin-bottom:10px; background:#fff;
}
.issue-card .cat { font-size:.72rem; font-weight:700; text-transform:uppercase; letter-spacing:.08em;}
.issue-card .orig { text-decoration:line-through; color:#B0341F; }
.issue-card .corr { color:#1D7A46; font-weight:600; }

/* ── Validation report (screenshot layout) ───── */
.vr-head { display:flex; gap:60px; margin:6px 0 20px; }
.vr-head .cell .lbl { color:#D71920; font-weight:700; font-size:1.15rem; }
.vr-head .cell .val { color:#2A2F3A; font-weight:600; font-size:1.05rem; margin-top:2px; }
.vr-sec-title { color:#D71920; font-weight:700; font-size:1.05rem; margin:4px 0 8px; }
.vr-box {
  border:3px solid #D71920; border-radius:2px; background:#fff;
  padding:12px 14px; min-height:90px; font-size:.9rem; color:#2A2F3A;
  white-space:pre-wrap; line-height:1.55;
}
.vr-err { color:#D71920; font-weight:700; font-size:.9rem; margin-top:8px; }
.vr-err .txt { font-weight:500; color:#8C1116; }
.vr-act { color:#1E9E3E; font-weight:700; font-size:.9rem; margin:2px 0 10px; }
.vr-act .txt { font-weight:500; color:#166D2C; }
.vr-ok-inline { color:#1E9E3E; font-weight:700; margin-top:8px; }
.vr-noerrors {
  display:inline-block; background:#1E9E3E; color:#fff; font-weight:700;
  font-size:1.1rem; padding:12px 34px; border-radius:10px; letter-spacing:.03em;
  box-shadow:0 3px 12px rgba(30,158,62,.35); margin:8px 0 14px;
}
.vr-panel {
  border:1px solid #E6E9EF; border-radius:14px; background:#fff;
  padding:16px 18px; box-shadow:0 2px 10px rgba(16,36,62,.05); margin-bottom:12px;
}
.vr-panel h5 { margin:0 0 8px; color:#10243E; }
.vr-panel .fld { font-size:.72rem; font-weight:700; text-transform:uppercase;
  letter-spacing:.08em; color:#6B7688; margin-top:10px; }
.vr-panel .txt { font-size:.86rem; color:#2A2F3A; white-space:pre-wrap; line-height:1.5;
  max-height:170px; overflow-y:auto; background:#F8FAFC; border:1px solid #EEF1F6;
  border-radius:8px; padding:8px 10px; margin-top:3px;}

/* ── Login card ──────────────────────────────── */
.login-card {
  max-width:420px; margin:6vh auto 0; background:#fff; border:1px solid #E6E9EF;
  border-radius:18px; padding:34px 36px; box-shadow:0 8px 30px rgba(16,36,62,.10);
}

/* ── Buttons & tabs polish ───────────────────── */
.stButton>button[kind="primary"] {
  background:linear-gradient(120deg,#1D4E89,#2E7D6B); border:none; border-radius:10px;
  font-weight:600;
}
.stTabs [data-baseweb="tab"] { font-weight:600; }
div[data-testid="stSidebar"] { background:#F6F8FB; }
</style>
""", unsafe_allow_html=True)


def hero():
    st.markdown(
        """
        <div class="slc-hero">
          <h1>🎓 SLC <span class="hl">Course Content</span> Checker</h1>
          <p>Compare live course pages · tracker sheet · official specifications — and proofread like a pro.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def stat(col, value, label, kind="info"):
    col.markdown(
        f'<div class="stat-card {kind}"><div class="num">{value}</div>'
        f'<div class="lbl">{label}</div></div>',
        unsafe_allow_html=True,
    )


def render_validation_html(course: dict, result: dict):
    """Render the on-page validation report in the agreed screenshot layout."""
    st.markdown(
        '<div class="vr-head">'
        f'<div class="cell"><div class="lbl">Level</div><div class="val">{html.escape(str(course.get("level") or "—"))}</div></div>'
        f'<div class="cell"><div class="lbl">Type</div><div class="val">{html.escape(str(course.get("course_type") or "—"))}</div></div>'
        f'<div class="cell"><div class="lbl">Number</div><div class="val">{html.escape(str(course.get("category_id") or "—"))}</div></div>'
        '</div>', unsafe_allow_html=True)

    if result.get("status") == "pass":
        st.markdown('<div class="vr-noerrors">✔ &nbsp;No Errors — course passed all validation checks</div>',
                    unsafe_allow_html=True)

    for sec in result.get("sections", []):
        left, right = st.columns([1, 1], gap="large")
        with left:
            st.markdown(f'<div class="vr-sec-title">Current {html.escape(sec["section"])}</div>',
                        unsafe_allow_html=True)
            st.markdown(f'<div class="vr-box">{html.escape((sec.get("current") or "(no content recorded)")[:1600])}</div>',
                        unsafe_allow_html=True)
        with right:
            st.markdown('<div style="height:34px"></div>', unsafe_allow_html=True)
            issues = sec.get("issues") or []
            if not issues:
                st.markdown('<div class="vr-ok-inline">✔ Correct — matches the specification</div>',
                            unsafe_allow_html=True)
            for n, iss in enumerate(issues, start=1):
                st.markdown(
                    f'<div class="vr-err">Errors identified {n:02d}: '
                    f'<span class="txt">{html.escape(str(iss.get("error","")))}</span></div>'
                    f'<div class="vr-act">Recommend Action {n:02d}: '
                    f'<span class="txt">{html.escape(str(iss.get("recommendation","")))}</span></div>',
                    unsafe_allow_html=True)
        st.write("")

    total = sum(len(s.get("issues") or []) for s in result.get("sections", []))
    st.divider()
    st.markdown("#### 📋 Summary of detected issues")
    if result.get("summary"):
        st.caption(result["summary"])
    if total == 0:
        st.success("No issues detected for this course. ✅")
    else:
        rows, k = [], 0
        for sec in result.get("sections", []):
            for iss in sec.get("issues") or []:
                k += 1
                rows.append({"#": k, "Section": sec["section"],
                             "Type": str(iss.get("type", "")).replace("_", " ").title(),
                             "Error identified": iss.get("error", ""),
                             "Recommended action": iss.get("recommendation", "")})
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)


# ═══════════════════════════════════════════════════════════════════
#  AUTHENTICATION GATE
# ═══════════════════════════════════════════════════════════════════

init_db()
ensure_default_users()

try:
    api_key = st.secrets.get("OPENROUTER_API_KEY", os.environ.get("OPENROUTER_API_KEY", ""))
except Exception:  # no secrets.toml present at all
    api_key = os.environ.get("OPENROUTER_API_KEY", "")
model = MODEL

if "auth" not in st.session_state:
    st.session_state.auth = None

if st.session_state.auth is None:
    hero()
    st.write("")
    _, mid, _ = st.columns([1, 1.1, 1])
    with mid:
        st.markdown("### 🔐 Sign in")
        with st.form("login_form"):
            u = st.text_input("Username", autocomplete="username")
            p = st.text_input("Password", type="password", autocomplete="current-password")
            ok = st.form_submit_button("Sign in", type="primary", use_container_width=True)
        if ok:
            user = verify_login(u, p)
            if user:
                st.session_state.auth = user
                st.rerun()
            else:
                st.error("Invalid username or password.")
        st.caption("Ask your administrator for an account. "
                   "Admins manage users on the 👥 Users page.")
    st.stop()

AUTH = st.session_state.auth
IS_ADMIN = AUTH["role"] == "admin"

hero()

with st.sidebar:
    st.markdown(f"### 👤 {AUTH['username']}")
    st.caption(("🛡️ Administrator" if IS_ADMIN else "🔎 User") +
               " · signed in")
    if st.button("🚪 Log out", use_container_width=True):
        st.session_state.auth = None
        st.session_state.pop("val_result", None)
        st.rerun()
    st.divider()
    st.markdown("### ⚙️ Settings")
    if api_key:
        st.success("🔑 OpenRouter key loaded from secrets")
    else:
        st.warning("No API key — AI checks are disabled; validation uses the "
                   "built-in comparison. Add `OPENROUTER_API_KEY` to "
                   "`.streamlit/secrets.toml` for full functionality.")


# ═══════════════════════════════════════════════════════════════════
#  PAGES
# ═══════════════════════════════════════════════════════════════════

def filter_bar(key_prefix: str, show_category: bool = True):
    """Level / Type (and optionally Category ID) filters — populated dynamically
    from the imported tracker sheet. Returns the filtered course list + selection."""
    opts = filter_options()
    if show_category:
        f1, f2, f3 = st.columns(3)
        cat = f1.selectbox("Category ID", ["All"] + opts["category_ids"], key=f"{key_prefix}_cat")
        lvl = f2.selectbox("Level", ["All"] + [str(l) for l in opts["levels"]], key=f"{key_prefix}_lvl")
        typ = f3.selectbox("Type", ["All"] + opts["types"], key=f"{key_prefix}_typ")
    else:
        f1, f2 = st.columns(2)
        cat = "All"
        lvl = f1.selectbox("Level", ["All"] + [str(l) for l in opts["levels"]], key=f"{key_prefix}_lvl")
        typ = f2.selectbox("Type", ["All"] + opts["types"], key=f"{key_prefix}_typ")
    matches = courses_filtered(cat, lvl, typ)
    return matches, {"category_id": cat, "level": lvl, "type": typ}


def validation_downloads(course: dict, result: dict, key_prefix: str):
    """Download Report buttons (PDF + Word) shown on the Run Check page."""
    safe = re.sub(r"[^A-Za-z0-9]+", "_", course.get("course_name", "course"))[:60]
    stamp = f"{datetime.now():%Y-%m-%d}"
    d1, d2 = st.columns(2)
    with d1:
        st.download_button(
            "⬇️ Download Report (PDF)",
            data=build_validation_pdf(course, result),
            file_name=f"Validation_Report_{safe}_{stamp}.pdf",
            mime="application/pdf", type="primary",
            use_container_width=True, key=f"{key_prefix}_pdf")
    with d2:
        st.download_button(
            "⬇️ Download Report (Word)",
            data=build_validation_docx(course, result),
            file_name=f"Validation_Report_{safe}_{stamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary", use_container_width=True, key=f"{key_prefix}_docx")


# ───────────────────────────────────────────────
# PAGE · IMPORT COURSES (admin)
# ───────────────────────────────────────────────
def page_import():
    st.subheader("Upload the Excel tracker sheet")
    st.caption("Expected columns: number/category ID, course name, course URL, spec URL, "
               "entry requirements, method of assessment, course overview. Level and Type "
               "are read from the sheet when present, otherwise derived from the course "
               "name — they feed the Run Check filters automatically.")

    up = st.file_uploader("Tracker sheet (.xlsx / .xls / .csv)", type=["xlsx", "xls", "csv"])

    if up:
        df = pd.read_csv(up) if up.name.lower().endswith(".csv") else pd.read_excel(up)
        df.columns = [str(c).strip() for c in df.columns]
        st.dataframe(df.head(10), use_container_width=True)

        def guess(*needles):
            for col in df.columns:
                low = col.lower()
                if any(n in low for n in needles):
                    return col
            return None

        cols = ["— none —"] + list(df.columns)

        def pick(label, guessed, key):
            idx = cols.index(guessed) if guessed in cols else 0
            return st.selectbox(label, cols, index=idx, key=key)

        st.markdown("#### Map your columns")
        c1, c2, c3 = st.columns(3)
        with c1:
            m_name = pick("Course name *", guess("course name", "title", "name"), "m1")
            m_url = pick("Course page URL", guess("course url", "page url", "link", "url"), "m2")
            m_cat = pick("Category ID / Number", guess("category", "number", "course id", "ref"), "m7")
        with c2:
            m_spec = pick("Specification URL", guess("spec"), "m3")
            m_entry = pick("Entry Requirements", guess("entry"), "m4")
            m_lvl = pick("Level (optional)", guess("level"), "m8")
        with c3:
            m_assess = pick("Method of Assessment", guess("assess"), "m5")
            m_over = pick("Course Overview", guess("overview", "description"), "m6")
            m_typ = pick("Type (optional)", guess("type"), "m9")
        st.caption("If Level / Type are not mapped they are derived automatically from "
                   "the course name (e.g. *Level 3 Diploma in Accounting*).")

        if st.button("📥 Import / update courses", type="primary", disabled=(m_name == "— none —")):
            def val(row, col):
                if col == "— none —":
                    return None
                v = row.get(col)
                return None if pd.isna(v) else str(v).strip()

            inserted = updated = skipped = 0
            for _, row in df.iterrows():
                name = val(row, m_name)
                if not name:
                    skipped += 1
                    continue
                result = upsert_course({
                    "course_name": name,
                    "course_url": val(row, m_url),
                    "spec_url": val(row, m_spec),
                    "entry_requirements": val(row, m_entry),
                    "method_of_assessment": val(row, m_assess),
                    "course_overview": val(row, m_over),
                    "category_id": val(row, m_cat),
                    "level": val(row, m_lvl),
                    "course_type": val(row, m_typ),
                })
                inserted += result == "inserted"
                updated += result == "updated"

            s1, s2, s3 = st.columns(3)
            stat(s1, inserted, "New courses", "ok")
            stat(s2, updated, "Updated", "info")
            stat(s3, skipped, "Skipped (no name)", "warn")

            # register every distinct specification document and link courses —
            # shared documents are detected so each is processed only once
            docs = sync_spec_documents()
            d1, d2, d3 = st.columns(3)
            stat(d1, docs["total"], "Unique spec documents", "info")
            stat(d2, docs["created"], "New documents registered", "ok")
            stat(d3, docs["total"] - docs["processed"], "Awaiting processing", "warn"
                 if docs["total"] - docs["processed"] else "ok")
            st.success("Tracker imported ✅ — the Run Check filters now reflect this "
                       "sheet. Process any new specification documents in "
                       "**📑 Spec Documents** (one-time extraction, reused for every check).")

    st.divider()
    if courses_now := all_courses():
        st.markdown("#### 📚 Courses in database")
        st.dataframe(
            pd.DataFrame(courses_now)[["id", "category_id", "level", "course_type",
                                       "course_name", "course_url", "spec_url", "updated_at"]],
            use_container_width=True, hide_index=True,
        )


# ───────────────────────────────────────────────
# PAGE · SPEC DOCUMENTS (admin)
# ───────────────────────────────────────────────
def page_spec_docs():
    st.subheader("Qualification specification documents")
    st.caption("Each **unique** specification document is read and AI-extracted **once** — "
               "the structured data (qualification name/level/type, entry requirements, "
               "assessment, requirements, learning outcomes, mandatory units) is stored in "
               "the database and **reused** by every course and every Run Check. Courses "
               "sharing the same document automatically share the same extraction. "
               "Reprocess only when a document has been updated.")

    if not all_courses():
        st.info("Import your tracker sheet first (📥 Import Courses).")
        return
    sync_spec_documents()  # keep documents/links in sync with the courses table
    docs = all_spec_docs()
    if not docs:
        st.warning("No specification document URLs found in the imported tracker.")
        return

    processed = [d for d in docs if d["status"] == "processed"]
    errored = [d for d in docs if d["status"] == "error"]
    pending = [d for d in docs if d["status"] not in ("processed", "error")]
    shared = [d for d in docs if (d.get("course_count") or 0) > 1]
    s1, s2, s3, s4 = st.columns(4)
    stat(s1, len(docs), "Unique documents", "info")
    stat(s2, len(processed), "Processed", "ok" if processed else "warn")
    stat(s3, len(pending) + len(errored), "Pending / error", "warn" if pending or errored else "ok")
    stat(s4, len(shared), "Shared by 2+ courses", "info")
    st.write("")

    # ── batch: process everything that hasn't been processed yet ──
    todo = pending + errored
    if todo:
        if st.button(f"⚙️ Process all unprocessed documents ({len(todo)})", type="primary",
                     key="proc_all"):
            bar = st.progress(0.0)
            line = st.empty()
            ok = err = 0
            for i, d in enumerate(todo, start=1):
                line.markdown(f"⏳ {i}/{len(todo)} — {d.get('spec_url') or d.get('filename')}")
                r = process_spec_document(d["id"], api_key, model)
                ok += r["status"] in ("processed", "reused", "unchanged")
                err += r["status"] == "error"
                bar.progress(i / len(todo))
            line.empty()
            st.success(f"Done — {ok} processed, {err} failed. Failed documents can be "
                       "fixed by uploading the file manually below.")
            st.rerun()
    else:
        st.success("All specification documents are processed ✅ — Run Check reuses the "
                   "stored data without re-reading any document.")

    # ── all documents, always visible — no picking required ──
    st.divider()
    st.markdown("#### 📄 All specification documents & stored data")
    st.caption("Every document and its stored extraction is listed below — Run Check "
               "validates against exactly what is saved here; no document re-reading "
               "at check time. Expand a document to view, edit or reprocess it.")
    q = st.text_input("🔎 Filter documents (URL / filename / qualification name / course)",
                      key="sd_q").strip().lower()

    shown = docs
    if q:
        def _hit(d):
            hay = " ".join([str(d.get("spec_url") or ""), str(d.get("filename") or ""),
                            str(d.get("extracted_json") or "")]).lower()
            if q in hay:
                return True
            return any(q in (x["course_name"] or "").lower() for x in courses_for_doc(d["id"]))
        shown = [d for d in docs if _hit(d)]
        st.caption(f"{len(shown)} of {len(docs)} document(s) match")

    ICON = {"processed": "✅", "error": "❌"}
    for d in shown:
        doc = d
        data = spec_doc_data(doc)
        title = (data.get("qualification_name") or doc.get("filename")
                 or doc.get("spec_url") or "document")
        label = (f"{ICON.get(doc['status'], '⏳')} {title[:95]} · "
                 f"{doc.get('course_count') or 0} course(s)")
        with st.expander(label):
            st.markdown(f"**URL:** {doc.get('spec_url') or '—'}  \n"
                        f"**Status:** `{doc.get('status')}`"
                        + (f" · extracted with **{doc.get('method')}**" if doc.get("method") else "")
                        + (f" · processed {doc.get('processed_at')}" if doc.get("processed_at") else "")
                        + (f"  \n**Error:** {doc.get('error')}" if doc.get("error") else ""))
            linked = courses_for_doc(doc["id"])
            if linked:
                st.markdown("**Courses:** " + " · ".join(
                    f"{x.get('category_id') or ''} {x['course_name']}".strip()
                    for x in linked[:10]) + (" …" if len(linked) > 10 else ""))

            # stored data — always shown
            if data:
                for key, lab in SPECDOC_FIELDS:
                    v = data.get(key)
                    if not v:
                        continue
                    if isinstance(v, list):
                        v = " • ".join(v)
                    st.markdown(f"**{lab}:** {v[:600]}{'…' if len(str(v)) > 600 else ''}")
            else:
                st.info("Not processed yet — no stored data.")

            b1, b2 = st.columns([1, 3])
            with b1:
                if st.button("⚙️ Process" if doc["status"] != "processed" else "🔄 Reprocess",
                             key=f"sd_proc_{doc['id']}", use_container_width=True):
                    with st.spinner("Reading & extracting …"):
                        r = process_spec_document(doc["id"], api_key, model,
                                                  force=(doc["status"] == "processed"))
                    (st.success if r["status"] != "error" else st.error)(r["detail"])
                    if r["status"] != "error":
                        st.rerun()
            with b2:
                edit = st.checkbox("✏️ Edit / upload replacement", key=f"sd_editok_{doc['id']}")

            if edit:
                up = st.file_uploader("Upload / replace the document file (.pdf / .docx)",
                                      type=["pdf", "docx"], key=f"sd_up_{doc['id']}")
                if up and st.button("📑 Process uploaded file", type="primary",
                                    key=f"sd_upgo_{doc['id']}"):
                    try:
                        text = read_uploaded_spec(up)
                    except Exception as e:
                        st.error(f"Could not read the file: {e}")
                        text = None
                    if text:
                        with st.spinner("Extracting …"):
                            r = process_spec_document(doc["id"], api_key, model, force=True,
                                                      uploaded_text=text, uploaded_name=up.name)
                        (st.success if r["status"] != "error" else st.error)(r["detail"])
                        if r["status"] != "error":
                            st.rerun()

                edited = {}
                g1, g2, g3 = st.columns(3)
                edited["qualification_name"] = g1.text_input(
                    "Qualification Name", data.get("qualification_name") or "",
                    key=f"sd_qn_{doc['id']}")
                edited["qualification_level"] = g2.text_input(
                    "Qualification Level", str(data.get("qualification_level") or ""),
                    key=f"sd_ql_{doc['id']}")
                edited["qualification_type"] = g3.text_input(
                    "Qualification Type", data.get("qualification_type") or "",
                    key=f"sd_qt_{doc['id']}")
                edited["entry_requirements"] = st.text_area(
                    "Entry Requirements", data.get("entry_requirements") or "", height=110,
                    key=f"sd_er_{doc['id']}")
                edited["method_of_assessment"] = st.text_area(
                    "Method of Assessment", data.get("method_of_assessment") or "", height=110,
                    key=f"sd_ma_{doc['id']}")
                edited["qualification_specification_requirements"] = st.text_area(
                    "Qualification Specification Requirements",
                    data.get("qualification_specification_requirements") or "", height=110,
                    key=f"sd_qr_{doc['id']}")
                lo = st.text_area("Learning Outcomes (one per line)",
                                  "\n".join(data.get("learning_outcomes") or []), height=90,
                                  key=f"sd_lo_{doc['id']}")
                mu = st.text_area("Mandatory Units (one per line)",
                                  "\n".join(data.get("mandatory_units") or []), height=90,
                                  key=f"sd_mu_{doc['id']}")
                edited["other_information"] = st.text_area(
                    "Other Relevant Information", data.get("other_information") or "",
                    height=80, key=f"sd_oi_{doc['id']}")
                if st.button("💾 Save extraction", type="primary", key=f"sd_save_{doc['id']}"):
                    edited["learning_outcomes"] = [l.strip() for l in lo.splitlines() if l.strip()]
                    edited["mandatory_units"] = [l.strip() for l in mu.splitlines() if l.strip()]
                    save_spec_doc(doc["id"],
                                  {"extracted_json": json.dumps(edited, ensure_ascii=False)})
                    propagate_doc_to_courses(doc["id"])
                    st.success("Extraction saved ✅ — reused by every linked course.")

    # ── compact status table ──
    with st.expander("📄 Status table (all documents)"):
        st.dataframe(pd.DataFrame([{
            "Status": {"processed": "✅ processed", "error": "❌ error"}.get(d["status"], "⏳ pending"),
            "Qualification": (spec_doc_data(d).get("qualification_name") or "—")[:80],
            "Document": (d.get("filename") or d.get("spec_url") or "—")[:100],
            "Courses": d.get("course_count") or 0,
            "Method": d.get("method") or "—",
            "Processed at": d.get("processed_at") or "—",
        } for d in docs]), use_container_width=True, hide_index=True)


# ───────────────────────────────────────────────
# PAGE · MANAGE COURSES (admin)
# ───────────────────────────────────────────────
def page_manage():
    st.subheader("Manage imported course records")
    courses = all_courses()
    if not courses:
        st.info("No courses in the database yet.")
        return

    matches, _ = filter_bar("mng")
    st.dataframe(pd.DataFrame(matches)[["id", "category_id", "level", "course_type",
                                        "course_name", "updated_at"]],
                 use_container_width=True, hide_index=True)
    if not matches:
        return

    name = st.selectbox("Course to edit / delete", [c["course_name"] for c in matches],
                        key="mng_course")
    course = next(c for c in matches if c["course_name"] == name)

    with st.expander("✏️ Edit course record", expanded=False):
        c1, c2, c3 = st.columns(3)
        n_cat = c1.text_input("Category ID / Number", course.get("category_id") or "",
                              key=f"mc_{course['id']}")
        n_lvl = c2.text_input("Level", course.get("level") or "", key=f"ml_{course['id']}")
        n_typ = c3.selectbox("Type", ["", *COURSE_TYPES],
                             index=(COURSE_TYPES.index(course["course_type"]) + 1
                                    if course.get("course_type") in COURSE_TYPES else 0),
                             key=f"mt_{course['id']}")
        n_url = st.text_input("Course URL", course.get("course_url") or "",
                              key=f"mu_{course['id']}")
        n_spec = st.text_input("Specification URL", course.get("spec_url") or "",
                               key=f"ms_{course['id']}")
        n_entry = st.text_area("Entry Requirements (tracker)",
                               course.get("entry_requirements") or "", height=110,
                               key=f"me_{course['id']}")
        n_assess = st.text_area("Method of Assessment (tracker)",
                                course.get("method_of_assessment") or "", height=110,
                                key=f"ma_{course['id']}")
        n_over = st.text_area("Course Overview (tracker)",
                              course.get("course_overview") or "", height=110,
                              key=f"mo_{course['id']}")
        if st.button("💾 Save changes", type="primary", key="mng_save"):
            update_course_fields(course["id"], {
                "category_id": n_cat or None, "level": n_lvl or None,
                "course_type": n_typ or None, "course_url": n_url or None,
                "spec_url": n_spec or None, "entry_requirements": n_entry or None,
                "method_of_assessment": n_assess or None,
                "course_overview": n_over or None,
            })
            st.success("Course updated ✅")
            st.rerun()

    with st.expander("🗑️ Delete course"):
        st.warning(f"This permanently removes **{course['course_name']}** and all of "
                   "its reports.")
        if st.checkbox("I understand", key="mng_del_ok"):
            if st.button("Delete this course", type="primary", key="mng_del"):
                delete_course(course["id"])
                st.success("Course deleted.")
                st.rerun()


# ───────────────────────────────────────────────
# PAGE · RUN CHECK (user + admin)
# ───────────────────────────────────────────────
def page_run_check():
    st.subheader("Run Check")
    st.caption("Validates the live course page against the **stored** qualification "
               "specification data — the specification document is **not** re-read or "
               "re-extracted here.")
    courses = all_courses()
    if not courses:
        st.info("No courses available yet — an administrator must import the tracker "
                "sheet first.")
        return

    # ── filters: Level · Type · Course (populated from the imported tracker) ──
    matches, _ = filter_bar("rc", show_category=False)
    if not matches:
        st.warning("No courses match the selected filters.")
        return
    def _lab(c):
        return f"{c.get('category_id') or '—'} — {c['course_name']}"
    pick = st.selectbox("Course", [_lab(c) for c in matches], key="rc_course")
    course = next(c for c in matches if _lab(c) == pick)
    doc = get_spec_doc(course.get("spec_doc_id"))
    spec_data = spec_doc_data(doc) if doc.get("status") == "processed" else {}

    # ── left: course details · right: stored specification data ──
    left, right = st.columns([1, 1], gap="large")
    with left:
        st.markdown(
            '<div class="vr-panel"><h5>📘 Course details (tracker)</h5>'
            f'<div class="fld">Number</div><div class="txt">{html.escape(str(course.get("category_id") or "—"))}</div>'
            f'<div class="fld">Level · Type</div><div class="txt">Level {html.escape(str(course.get("level") or "—"))} · {html.escape(str(course.get("course_type") or "—"))}</div>'
            f'<div class="fld">Course page</div><div class="txt">{html.escape(str(course.get("course_url") or "—"))}</div>'
            f'<div class="fld">Entry Requirements (record)</div><div class="txt">{html.escape((course.get("entry_requirements") or "—")[:900])}</div>'
            f'<div class="fld">Method of Assessment (record)</div><div class="txt">{html.escape((course.get("method_of_assessment") or "—")[:900])}</div>'
            '</div>', unsafe_allow_html=True)
    with right:
        if spec_data:
            lo_n = len(spec_data.get("learning_outcomes") or [])
            mu_n = len(spec_data.get("mandatory_units") or [])
            st.markdown(
                '<div class="vr-panel"><h5>📑 Stored specification data</h5>'
                f'<div class="fld">Document</div><div class="txt">{html.escape(str(doc.get("filename") or doc.get("spec_url") or "—"))} · processed {html.escape(str(doc.get("processed_at") or ""))}</div>'
                f'<div class="fld">Qualification</div><div class="txt">{html.escape(str(spec_data.get("qualification_name") or "—"))} · Level {html.escape(str(spec_data.get("qualification_level") or "—"))} · {html.escape(str(spec_data.get("qualification_type") or "—"))}</div>'
                f'<div class="fld">Entry Requirements</div><div class="txt">{html.escape((spec_data.get("entry_requirements") or "—")[:700])}</div>'
                f'<div class="fld">Method of Assessment</div><div class="txt">{html.escape((spec_data.get("method_of_assessment") or "—")[:700])}</div>'
                f'<div class="fld">Specification Requirements</div><div class="txt">{html.escape((spec_data.get("qualification_specification_requirements") or "—")[:700])}</div>'
                f'<div class="fld">Learning Outcomes · Mandatory Units</div><div class="txt">{lo_n} outcome(s) · {mu_n} unit(s) stored</div>'
                '</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="vr-panel"><h5>📑 Stored specification data</h5>'
                        '<div class="txt">— not processed yet —</div></div>',
                        unsafe_allow_html=True)
            st.warning("The specification document for this course has not been processed "
                       "yet. An administrator needs to process it once in **📑 Spec "
                       "Documents**; after that every check reuses the stored data.")

    # ── run the validation ──
    st.write("")
    if st.button("🔍 Run Check", type="primary", key="rc_validate",
                 disabled=not spec_data):
        page_text = ""
        if course.get("course_url"):
            with st.spinner("Loading the live course page …"):
                try:
                    page_text = extract_page_text(course["course_url"])
                except Exception as e:
                    st.warning(f"Could not load the course page ({e}) — validating the "
                               "tracker record instead.")
        with st.spinner(f"Validating {course['course_name']} against the stored "
                        "specification data …"):
            result = validate_course_vs_spec(course, spec_data, page_text, api_key, model)
        save_validation_report(course["id"], AUTH["username"], result["status"],
                               result.get("summary", ""), result)
        st.session_state["val_result"] = {"course_id": course["id"], "result": result}

    cached = st.session_state.get("val_result")
    if cached and cached["course_id"] == course["id"]:
        result = cached["result"]
        st.write("")
        render_validation_html(course, result)
        st.divider()
        st.markdown("#### ⬇️ Download Report")
        validation_downloads(course, result, "rc")
    else:
        prior = [r for r in latest_validation_reports() if r["course_id"] == course["id"]]
        if prior:
            r = prior[0]
            st.caption(f"Last validated {r['checked_at']} by {r['checked_by']} — "
                       f"status: {'✅ pass' if r['status'] == 'pass' else '⚠️ errors'}. "
                       "Run Check to refresh.")


# ───────────────────────────────────────────────
# PAGE · GRAMMAR CHECK (its own tab — moved out of Run Check)
# ───────────────────────────────────────────────
def page_grammar():
    st.subheader("Grammar Check")
    st.caption("Paste or upload course content. The reviewer highlights grammar, "
               "articles, sentence structure, capitalisation, proper nouns, spelling, "
               "commas and punctuation consistency — colour-coded like a "
               "proofreader's markup.")

    legend = "".join(
        f'<span style="background:{v["bg"]};border:1px solid {v["border"]};color:#2A2F3A;">{v["label"]}</span>'
        for v in QR_CATEGORIES.values()
    )
    st.markdown(f'<div class="qr-legend">{legend}</div>', unsafe_allow_html=True)
    st.write("")

    src = st.radio("Input", ["Paste text", "Upload file (.txt / .docx)",
                             "Use a course's overview"], horizontal=True, key="gc_src")
    text = ""
    if src == "Paste text":
        text = st.text_area("Course content to review", height=220, key="gc_paste",
                            placeholder="Paste the course description, overview or any page copy here…")
    elif src == "Use a course's overview":
        names = [c["course_name"] for c in all_courses()]
        if not names:
            st.info("No courses imported yet.")
        else:
            cname = st.selectbox("Course", names, key="gc_course")
            crs = next(c for c in all_courses() if c["course_name"] == cname)
            text = crs.get("course_overview") or ""
            if text:
                st.text_area("Loaded content", text, height=180, key="qr_course_text")
            else:
                st.info("This course has no overview text in the tracker.")
    else:
        f = st.file_uploader("Upload content", type=["txt", "docx"], key="qr_up")
        if f:
            if f.name.lower().endswith(".docx"):
                d = Document(io.BytesIO(f.read()))
                text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
            else:
                text = f.read().decode("utf-8", errors="ignore")
            st.text_area("Loaded content", text, height=180)

    if st.button("✍️ Review content quality", type="primary", disabled=not text.strip(),
                 key="gc_go"):
        if not api_key:
            st.error("No API key found — add OPENROUTER_API_KEY to .streamlit/secrets.toml and restart.")
        else:
            with st.spinner("Proofreading …"):
                try:
                    result = run_quality_review(text, api_key, model)
                    st.session_state["qr_result"] = result
                    st.session_state["qr_text"] = text
                except Exception as e:
                    st.error(f"Review failed: {e}")

    if "qr_result" in st.session_state:
        result = st.session_state["qr_result"]
        text = st.session_state["qr_text"]
        issues = result.get("issues", [])

        c1, c2, c3 = st.columns(3)
        stat(c1, len(issues), "Issues found", "err" if issues else "ok")
        top_cat = max({i.get("category") for i in issues},
                      key=lambda c: sum(1 for i in issues if i.get("category") == c),
                      default="—")
        stat(c2, QR_CATEGORIES.get(top_cat, {}).get("label", "—"), "Most common issue", "warn")
        stat(c3, f"{max(0, 100 - len(issues) * 4)}%", "Quality score", "info")
        st.write("")

        view_marked, view_fixed, view_list = st.tabs(
            ["🖍️ Marked-up text", "✅ Corrected text", "📋 Issue list"])

        with view_marked:
            if issues:
                st.markdown(f'<div class="qr-paper">{annotate_text_html(text, issues)}</div>',
                            unsafe_allow_html=True)
                st.caption("Hover a highlight to see the correction and explanation.")
            else:
                st.success("No issues found — this content is clean. 🎉")

        with view_fixed:
            corrected = result.get("corrected_text", text)
            st.text_area("Corrected version (copy-ready)", corrected, height=260)
            st.download_button("⬇️ Download corrected text", corrected,
                               file_name="corrected_content.txt")

        with view_list:
            if not issues:
                st.success("Nothing to list — no issues found.")
            for n, issue in enumerate(issues, start=1):
                cat = issue.get("category", "grammar")
                sty = QR_CATEGORIES.get(cat, QR_CATEGORIES["grammar"])
                st.markdown(
                    f'<div class="issue-card" style="border-left-color:{sty["border"]}">'
                    f'<div class="cat" style="color:{sty["border"]}">#{n} · {sty["label"]}</div>'
                    f'<span class="orig">{html.escape(str(issue.get("original","")))}</span> → '
                    f'<span class="corr">{html.escape(str(issue.get("correction","")))}</span><br>'
                    f'<small>{html.escape(str(issue.get("explanation","")))}</small>'
                    f'</div>', unsafe_allow_html=True)


# ───────────────────────────────────────────────
# PAGE · OTHER CHECKS (existing live-page 3-way audit — its own tab)
# ───────────────────────────────────────────────
def page_other_checks():
    st.subheader("Other Checks")
    st.caption("Checks the live course page against the tracker and the official "
               "specification — the original 3-way audit.")
    courses = all_courses()
    if not courses:
        st.info("No courses available yet.")
        return

    matches, _ = filter_bar("oc")
    if not matches:
        st.warning("No courses match the selected filters.")
        return
    name = st.selectbox("Course", [c["course_name"] for c in matches], key="oc_course")
    course = next(c for c in matches if c["course_name"] == name)

    selected_fields = [k for k, label in FIELD_OPTIONS.items()
                       if st.checkbox(label, value=True, key=f"oc_field_{k}")]
    if not selected_fields:
        st.warning("Select at least one field to check.")
        return
    st.caption(f"Checking fields: **{', '.join(FIELD_OPTIONS[k] for k in selected_fields)}**")

    colA, colB = st.columns([1, 1], gap="large")

    # ── Single course check ──
    with colA:
        st.markdown("#### 🎯 Check this course")
        if st.button("🔍 Check this course", type="primary", key="single"):
            if not api_key:
                st.error("No API key found — add OPENROUTER_API_KEY to .streamlit/secrets.toml and restart.")
            else:
                with st.spinner(f"Checking {name} …"):
                    rep = check_course(course, selected_fields, api_key, model)
                save_report(course["id"], selected_fields, rep["status"],
                            rep["errors"], rep["summary"], rep.get("rewrites"))
                if rep["status"] == "pass":
                    st.success(f"✅ **{name}** — no errors found. {rep['summary']}")
                elif rep["status"] == "errors":
                    st.error(f"⚠️ **{name}** — {len(rep['errors'])} issue(s). {rep['summary']}")
                    for err in rep["errors"]:
                        with st.expander(f"❌ {err.get('field')} · {str(err.get('severity','')).upper()} — {err.get('issue','')[:70]}"):
                            st.markdown(f"**Issue:** {err.get('issue')}")
                            st.markdown(f"**Live page says:** {err.get('live_content')}")
                            st.markdown(f"**Expected:** {err.get('expected_content')}")
                            st.markdown(f"**💡 Suggested fix:** {err.get('suggested_fix')}")
                    for rw in rep.get("rewrites", []):
                        orig = rw.get("originality", 0)
                        badge = "🟢" if orig >= 90 else ("🟡" if orig >= 75 else "🔴")
                        with st.expander(f"✨ Suggested wording — {rw.get('field')} "
                                         f"({badge} {orig}% original)", expanded=True):
                            st.text_area("Copy-ready replacement text",
                                         rw.get("suggested_wording", ""),
                                         height=160, key=f"rw_{name}_{rw.get('field')}")
                            st.caption(f"Originality {orig}% — share of 4-word phrases NOT "
                                       "found in the specification, live page or tracker. "
                                       "Verified plagiarism-safe before display.")
                else:
                    st.warning(f"Check failed — {rep['summary']}")

    # ── Bulk check ──
    with colB:
        st.markdown("#### 🚀 Bulk check")
        targets = matches
        st.caption(f"{len(targets)} filtered course(s) will be checked")
        workers = st.slider("Parallel workers", min_value=1, max_value=10, value=6, key="bulk_workers",
                            help="How many courses to check at the same time. Higher = faster, "
                                 "but lower it if you hit API rate limits.")
        if st.button(f"🚀 Run bulk check on {len(targets)} courses", type="primary", key="bulk"):
            if not api_key:
                st.error("No API key found — add OPENROUTER_API_KEY to .streamlit/secrets.toml and restart.")
            else:
                bar = st.progress(0.0)
                status = st.empty()
                results = {"pass": 0, "errors": 0, "failed": 0}
                log = st.container()
                done = 0
                t0 = time.time()
                # Workers only do network + LLM work (thread-safe).
                # DB writes and UI updates happen here in the main thread.
                with ThreadPoolExecutor(max_workers=workers) as pool:
                    futures = {
                        pool.submit(check_course, c, selected_fields, api_key, model): c
                        for c in targets
                    }
                    for fut in as_completed(futures):
                        cc = futures[fut]
                        try:
                            rep = fut.result()
                        except Exception as e:
                            rep = {"course_id": cc["id"], "course_name": cc["course_name"],
                                   "status": "failed", "summary": f"Unexpected error: {e}",
                                   "errors": [], "rewrites": []}
                        save_report(cc["id"], selected_fields, rep["status"],
                                    rep["errors"], rep["summary"], rep.get("rewrites"))
                        results[rep["status"]] += 1
                        done += 1
                        icon = {"pass": "✅", "errors": "⚠️", "failed": "❌"}[rep["status"]]
                        log.markdown(f"{icon} **{cc['course_name']}** — "
                                     f"{len(rep['errors'])} issue(s). {rep['summary'][:120]}")
                        status.markdown(f"⏳ **{done}/{len(targets)}** done "
                                        f"({time.time() - t0:.0f}s elapsed)")
                        bar.progress(done / len(targets))
                status.empty()
                r1, r2, r3 = st.columns(3)
                stat(r1, results["pass"], "Passed", "ok")
                stat(r2, results["errors"], "With errors", "err")
                stat(r3, results["failed"], "Failed to check", "warn")
                st.success("Bulk check complete — reports saved to the database. "
                           "Head to 📊 Reports to export the Word document.")


# ───────────────────────────────────────────────
# PAGE · REPORTS (user: view/download · admin: full)
# ───────────────────────────────────────────────
def page_reports():
    st.subheader("Reports")

    tab_val, tab_audit = st.tabs(["🧾 Validation reports", "📊 Live-page audit reports"])

    # ── validation reports (new) ──
    with tab_val:
        vreports = latest_validation_reports()
        if not vreports:
            st.info("No validation reports yet — run a validation on the 🔍 Run Check page.")
        else:
            n_pass = sum(r["status"] == "pass" for r in vreports)
            n_err = sum(r["status"] == "errors" for r in vreports)
            c1, c2, c3 = st.columns(3)
            stat(c1, len(vreports), "Courses validated", "info")
            stat(c2, n_pass, "No errors", "ok")
            stat(c3, n_err, "With errors", "err")
            st.write("")

            st.dataframe(pd.DataFrame([{
                "Number": r.get("category_id"), "Level": r.get("level"),
                "Type": r.get("course_type"), "Course": r["course_name"],
                "Status": "🟢 No Errors" if r["status"] == "pass" else "⚠️ Errors",
                "Issues": sum(len(s.get("issues") or [])
                              for s in json.loads(r["results_json"] or "{}").get("sections", [])),
                "Checked at": r["checked_at"], "Checked by": r.get("checked_by"),
            } for r in vreports]), use_container_width=True, hide_index=True)

            sel = st.selectbox("View / download a validation report",
                               [r["course_name"] for r in vreports], key="vr_pick")
            r = next(x for x in vreports if x["course_name"] == sel)
            course = get_course(r["course_id"])
            result = json.loads(r["results_json"] or "{}")
            with st.expander("🔎 View report", expanded=False):
                render_validation_html(course, result)
            validation_downloads(course, result, "rep")

    # ── existing live-page audit reports + Word export (unchanged) ──
    with tab_audit:
        reports = latest_reports()

        if not reports:
            st.info("No reports yet — run some checks first.")
        else:
            n_pass = sum(r["status"] == "pass" for r in reports)
            n_err = sum(r["status"] == "errors" for r in reports)
            n_fail = sum(r["status"] == "failed" for r in reports)
            c1, c2, c3, c4 = st.columns(4)
            stat(c1, len(reports), "Courses checked", "info")
            stat(c2, n_pass, "Passed", "ok")
            stat(c3, n_err, "With errors", "err")
            stat(c4, n_fail, "Check failed", "warn")
            st.write("")

            table = pd.DataFrame([{
                "Course": r["course_name"],
                "Status": {"pass": "✅ Pass", "errors": "⚠️ Errors", "failed": "❌ Failed"}[r["status"]],
                "Issues": len(json.loads(r["errors_json"] or "[]")),
                "Checked at": r["checked_at"],
                "Summary": r["summary"],
            } for r in reports])
            st.dataframe(table, use_container_width=True, hide_index=True)

            # detail viewer
            err_reports_db = [r for r in reports if r["status"] == "errors"]
            if err_reports_db:
                with st.expander("🔎 View error details"):
                    sel = st.selectbox("Course with errors", [r["course_name"] for r in err_reports_db])
                    r = next(x for x in err_reports_db if x["course_name"] == sel)
                    for err in json.loads(r["errors_json"]):
                        style = {"high": "#E5484D", "medium": "#F76B15", "low": "#30A46C"}.get(
                            str(err.get("severity", "medium")).lower(), "#F76B15")
                        st.markdown(
                            f'<div class="issue-card" style="border-left-color:{style}">'
                            f'<div class="cat" style="color:{style}">{err.get("field","")} · {err.get("severity","")}</div>'
                            f'<b>{html.escape(str(err.get("issue","")))}</b><br>'
                            f'<span class="orig">{html.escape(str(err.get("live_content","")))}</span> → '
                            f'<span class="corr">{html.escape(str(err.get("suggested_fix","")))}</span>'
                            f'</div>', unsafe_allow_html=True)
                    for rw in json.loads(r.get("rewrites_json") or "[]"):
                        orig = rw.get("originality", 0)
                        badge = "🟢" if orig >= 90 else ("🟡" if orig >= 75 else "🔴")
                        st.markdown(f"**✨ Suggested wording — {rw.get('field')} "
                                    f"({badge} {orig}% original)**")
                        st.text_area("Copy-ready replacement text", rw.get("suggested_wording", ""),
                                     height=150, key=f"rep_rw_{r['course_id']}_{rw.get('field')}")

            st.divider()
            st.markdown("#### 📝 Word document — courses with errors only")
            if not err_reports_db:
                st.success("No courses with errors — nothing to export. 🎉")
            else:
                word_input = [{
                    "course_name": r["course_name"],
                    "course_url": r["course_url"],
                    "summary": r["summary"],
                    "errors": json.loads(r["errors_json"] or "[]"),
                    "rewrites": json.loads(r.get("rewrites_json") or "[]"),
                } for r in err_reports_db]
                docx_bytes = build_word_report(word_input)
                st.download_button(
                    f"⬇️ Download Word report ({len(err_reports_db)} courses with errors)",
                    data=docx_bytes,
                    file_name=f"SLC_Course_Error_Report_{datetime.now():%Y-%m-%d}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                )


# ───────────────────────────────────────────────
# PAGE · USERS (admin)
# ───────────────────────────────────────────────
def page_users():
    st.subheader("User accounts")
    users = all_users()
    admins = [u for u in users if u["role"] == "admin"]

    st.dataframe(pd.DataFrame([{
        "Username": u["username"],
        "Role": "🛡️ Admin" if u["role"] == "admin" else "🔎 User",
        "Created": u["created_at"],
    } for u in users]), use_container_width=True, hide_index=True)

    c1, c2 = st.columns(2, gap="large")

    with c1:
        st.markdown("#### ➕ Add account")
        with st.form("add_user", clear_on_submit=True):
            nu = st.text_input("Username")
            np_ = st.text_input("Password", type="password")
            nr = st.selectbox("Role", ["user", "admin"])
            ok = st.form_submit_button("Create account", type="primary")
        if ok:
            if not nu.strip() or len(np_) < 6:
                st.error("Username required and password must be at least 6 characters.")
            elif create_user(nu, np_, nr):
                st.success(f"Account **{nu}** created ✅")
                st.rerun()
            else:
                st.error("That username already exists.")

    with c2:
        st.markdown("#### 🔧 Reset password / delete")
        target_name = st.selectbox("Account", [u["username"] for u in users], key="u_target")
        target = next(u for u in users if u["username"] == target_name)
        newpw = st.text_input("New password", type="password", key="u_newpw")
        if st.button("Reset password", disabled=len(newpw) < 6):
            set_password(target["id"], newpw)
            st.success("Password updated ✅")
        st.write("")
        is_last_admin = target["role"] == "admin" and len(admins) <= 1
        is_self = target["username"] == AUTH["username"]
        if is_last_admin:
            st.caption("⚠️ This is the last admin account — it can't be deleted.")
        elif is_self:
            st.caption("⚠️ You can't delete the account you are signed in with.")
        elif st.button("🗑️ Delete this account"):
            delete_user(target["id"])
            st.success("Account deleted.")
            st.rerun()

    st.divider()
    st.caption("Default accounts on first run — **admin / admin123** and "
               "**user / user123**. Change these passwords before sharing the tool.")


# ═══════════════════════════════════════════════════════════════════
#  NAVIGATION (role-based)
# ═══════════════════════════════════════════════════════════════════

if IS_ADMIN:
    t_imp, t_spec, t_mng, t_run, t_gram, t_oth, t_rep, t_usr = st.tabs(
        ["📥 Import Courses", "📑 Spec Documents", "🗂 Manage Courses",
         "🔍 Run Check", "✍️ Grammar Check", "🧪 Other Checks",
         "📊 Reports", "👥 Users"])
    with t_imp:
        page_import()
    with t_spec:
        page_spec_docs()
    with t_mng:
        page_manage()
    with t_run:
        page_run_check()
    with t_gram:
        page_grammar()
    with t_oth:
        page_other_checks()
    with t_rep:
        page_reports()
    with t_usr:
        page_users()
else:
    # Users get Run Check (validate, view & download reports), Grammar Check,
    # Other Checks and Reports. No upload or data-modification pages.
    t_run, t_gram, t_oth, t_rep = st.tabs(
        ["🔍 Run Check", "✍️ Grammar Check", "🧪 Other Checks", "📊 Reports"])
    with t_run:
        page_run_check()
    with t_gram:
        page_grammar()
    with t_oth:
        page_other_checks()
    with t_rep:
        page_reports()
