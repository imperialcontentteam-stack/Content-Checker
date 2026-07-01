import io
import json
import pandas as pd
from docx import Document
from docx.shared import Inches
from modules.database import fetch_all, fetch_one


def reports_to_dataframe():
    rows = fetch_all(
        """
        select id, course_name, checked_fields, summary_status, result_json, created_at
        from check_reports
        order by created_at desc
        """
    )

    flattened = []
    for row in rows:
        try:
            result = json.loads(row.get("result_json") or "[]")
        except Exception:
            result = []

        if isinstance(result, dict):
            result = [result]

        if not result:
            flattened.append(row)
            continue

        for item in result:
            flattened.append(
                {
                    "Report ID": row["id"],
                    "Course Name": row["course_name"],
                    "Created At": row["created_at"],
                    "Summary Status": row["summary_status"],
                    "Field Checked": item.get("field_checked"),
                    "Decision": item.get("decision"),
                    "Priority": item.get("priority"),
                    "Website Evidence": item.get("current_website_evidence"),
                    "Tracker Evidence": item.get("tracker_sheet_evidence"),
                    "Specification Evidence": item.get("specification_evidence"),
                    "Explanation": item.get("explanation"),
                    "Suggested Corrected Wording": item.get("suggested_corrected_wording"),
                    "Suggested Action": item.get("suggested_action"),
                    "Similarity Risk": item.get("wording_similarity_risk"),
                    "Course Overview Plagiarism %": item.get("plagiarism_percentage") if item.get("field_checked") == "Course Overview" else "",
                    "Low Risk Rewritten Wording": item.get("low_risk_rewritten_wording"),
                }
            )

    return pd.DataFrame(flattened)


def dataframe_to_excel_bytes(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Reports", index=False)
        worksheet = writer.sheets["Reports"]
        for idx, col in enumerate(df.columns):
            width = min(max(len(str(col)) + 2, 15), 60)
            worksheet.set_column(idx, idx, width)
    output.seek(0)
    return output.getvalue()


def _add_key_value(paragraph_parent, label, value):
    p = paragraph_parent.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value or ""))


def _safe_json_list(value):
    try:
        parsed = json.loads(value or "[]")
    except Exception:
        parsed = []
    if isinstance(parsed, dict):
        return [parsed]
    return parsed if isinstance(parsed, list) else []


def build_single_report_docx_bytes(report_id: int):
    row = fetch_one(
        """
        select id, course_name, checked_fields, summary_status, result_json, created_at
        from check_reports
        where id = ?
        """,
        (report_id,),
    )
    if not row:
        raise ValueError("Report not found.")

    results = _safe_json_list(row.get("result_json"))
    return report_json_to_docx_bytes(
        course_name=row.get("course_name"),
        summary_status=row.get("summary_status"),
        results=results,
        report_id=row.get("id"),
        created_at=row.get("created_at"),
    )


def report_json_to_docx_bytes(course_name, summary_status, results, report_id=None, created_at=None):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.add_heading("SLC Course Content Check Report", level=1)
    _add_key_value(doc, "Course", course_name)
    if report_id:
        _add_key_value(doc, "Report ID", report_id)
    if created_at:
        _add_key_value(doc, "Created At", created_at)
    _add_key_value(doc, "Summary Status", summary_status)

    for item in results:
        doc.add_heading(item.get("field_checked") or "Checked Field", level=2)
        _add_key_value(doc, "Decision", item.get("decision"))
        _add_key_value(doc, "Priority", item.get("priority"))
        _add_key_value(doc, "Similarity Risk", item.get("wording_similarity_risk"))
        if item.get("field_checked") == "Course Overview":
            plagiarism_value = item.get("plagiarism_percentage")
            _add_key_value(doc, "Course Overview Plagiarism %", "N/A" if plagiarism_value is None else f"{plagiarism_value}%")

        doc.add_heading("Explanation", level=3)
        doc.add_paragraph(item.get("explanation") or "")

        doc.add_heading("Suggested corrected wording", level=3)
        doc.add_paragraph(item.get("suggested_corrected_wording") or "")

        doc.add_heading("Low-risk rewritten wording", level=3)
        doc.add_paragraph(item.get("low_risk_rewritten_wording") or "")

        doc.add_heading("Evidence", level=3)
        _add_key_value(doc, "Website evidence", item.get("current_website_evidence"))
        _add_key_value(doc, "Tracker evidence", item.get("tracker_sheet_evidence"))
        _add_key_value(doc, "Specification evidence", item.get("specification_evidence"))
        _add_key_value(doc, "Suggested action", item.get("suggested_action"))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def all_reports_to_docx_bytes(df: pd.DataFrame):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.add_heading("SLC Course Content Check Reports", level=1)
    doc.add_paragraph(f"Total report rows: {len(df)}")

    if df.empty:
        doc.add_paragraph("No report data available.")
    else:
        for _, row in df.iterrows():
            doc.add_heading(str(row.get("Course Name") or "Course"), level=2)
            _add_key_value(doc, "Report ID", row.get("Report ID"))
            _add_key_value(doc, "Created At", row.get("Created At"))
            _add_key_value(doc, "Field Checked", row.get("Field Checked"))
            _add_key_value(doc, "Decision", row.get("Decision"))
            _add_key_value(doc, "Priority", row.get("Priority"))
            if row.get("Field Checked") == "Course Overview":
                _add_key_value(doc, "Course Overview Plagiarism %", row.get("Course Overview Plagiarism %"))
            _add_key_value(doc, "Summary Status", row.get("Summary Status"))
            _add_key_value(doc, "Explanation", row.get("Explanation"))
            _add_key_value(doc, "Suggested Corrected Wording", row.get("Suggested Corrected Wording"))
            _add_key_value(doc, "Suggested Action", row.get("Suggested Action"))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
